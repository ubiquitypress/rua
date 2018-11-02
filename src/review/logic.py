import json
import mimetypes as mime
import os
from os import path
from uuid import uuid4

from django.conf import settings
from django.contrib import messages
from django.core.files.storage import default_storage
from django.http import StreamingHttpResponse, HttpResponseRedirect
from django.shortcuts import get_object_or_404
from django.utils.encoding import smart_text

from bs4 import BeautifulSoup
from docx import Document

from core import models as core_models, logic as core_logic
from core.decorators import is_reviewer
from review import models


def notify_editors(
        book,
        message,
        book_editors,
        series_editor,
        creator,
        workflow
):
    for editor in book_editors:
        notification = core_models.Task(
            book=book,
            assignee=editor,
            creator=creator,
            text=message,
            workflow=workflow,
        )
        notification.save()

    if series_editor:
        notification = core_models.Task(
            book=book,
            assignee=series_editor,
            creator=creator,
            text=message,
            workflow=workflow,
        )
        notification.save()


def has_additional_files(submission):
    additional_files = [
        _file for _file in submission.files.all() if _file.kind == 'additional'
    ]

    return True if additional_files else False


def handle_review_file(_file, content_type, review_assignment, kind):
    original_filename = smart_text(
        _file._get_name().replace(',', '_').replace(';', '_')
    )
    filename = str(uuid4()) + str(path.splitext(original_filename)[1])
    file_mime = mime.guess_type(filename)[0] or 'application/octet-stream'

    if content_type == 'book':
        file_path = path.join(
            settings.BOOK_DIR,
            str(review_assignment.book.id),
            str(filename),
        )
    else:
        file_path = path.join(
            settings.PROPOSAL_DIR,
            str(review_assignment.proposal.id),
            str(filename),
        )

    with default_storage.open(file_path, 'wb') as file_stream:
        file_stream.write(_file.read())

    new_file = core_models.File.objects.create(
        mime_type=file_mime,
        original_filename=original_filename,
        uuid_filename=filename,
        stage_uploaded=1,
        kind=kind,
        owner=review_assignment.user,
    )
    review_assignment.files.add(new_file)

    return file_path


def create_review_form(submission, review_form):
    document = Document()
    document.add_heading(submission.title, 0)
    p = document.add_paragraph(
        'You should complete this form and then use the review '
        'page to upload it.'
    )
    relations = models.FormElementsRelationship.objects.filter(
        form=review_form
    ).order_by('order')

    for relation in relations:
        if relation.element.field_type in ['text', 'textarea', 'date', 'email']:
            document.add_heading(
                relation.element.name + ": _______________________________",
                level=1,
            )
            document.add_paragraph(relation.help_text).italic = True

        if relation.element.field_type in ['select', 'check']:
            document.add_heading(relation.element.name, level=1)

            if relation.element.field_type == 'select':
                choices = render_choices(relation.element.choices)
            else:
                choices = ['Y', 'N']

            p = document.add_paragraph(relation.help_text)
            p.add_run(
                ' Mark your choice however you like, as long as it is clear.'
            ).italic = True
            table = document.add_table(rows=2, cols=len(choices))
            hdr_cells = table.rows[0].cells

            for i, choice in enumerate(choices):
                hdr_cells[i].text = choice[0]

            table.style = 'TableGrid'

    document.add_page_break()

    if not os.path.exists(os.path.join(settings.BASE_DIR, 'files', 'forms')):
        os.makedirs(os.path.join(settings.BASE_DIR, 'files', 'forms'))

    path = os.path.join(
        settings.BASE_DIR,
        'files',
        'forms',
        '%s.docx' % str(uuid4())
    )

    document.save(path)

    return path


@is_reviewer
def serve_file(request, file_path):
    try:
        file_stream = default_storage.open(file_path, 'r')
        mimetype = mime.guess_type(file_path)
        response = StreamingHttpResponse(file_stream, content_type=mimetype)
        response['Content-Disposition'] = (
            "attachment; filename=review_form.docx"
        )
        return response
    except IOError:
        messages.add_message(request, messages.ERROR, 'File not found.')
        return HttpResponseRedirect(request.META.get('HTTP_REFERER'))


def handle_editorial_review_file(
        file,
        proposal,
        review_assignment,
        kind,
        editorial
):
    original_filename = smart_text(file._get_name())
    filename = str(uuid4()) + str(os.path.splitext(original_filename)[1])
    file_mime = mime.guess_type(filename)[0] or 'application/octet-stream'
    file_path = os.path.join(
        settings.BOOK_DIR,
        str(review_assignment.book.id),
        filename
    )

    with default_storage.open(file_path, 'wb') as file_stream:
        file_stream.write(file.read())

    new_file = core_models.File.objects.create(
        mime_type=file_mime,
        original_filename=original_filename,
        uuid_filename=filename,
        stage_uploaded=1,
        kind=kind,
        owner=review_assignment.management_editor,
    )

    if editorial:
        review_assignment.editorial_board_files.add(new_file)
    else:
        review_assignment.publication_committee_files.add(new_file)

    return path


def render_choices(choices):
    c_split = choices.split('|')
    return [(choice.capitalize(), choice) for choice in c_split]


def create_completed_review_form(submission, review_id):
    document = Document()
    document.add_heading(submission.title, 0)
    review_assignment = get_object_or_404(
        core_models.ReviewAssignment,
        pk=review_id,
    )
    if review_assignment.review_form:
        relations = models.FormElementsRelationship.objects.filter(
            form=review_assignment.review_form,
        ).order_by('order')
    else:
        review_assignment.review_form = submission.review_form
        review_assignment.save()
        relations = models.FormElementsRelationship.objects.filter(
            form=submission.review_form,
        ).order_by('order')

    if review_assignment.results:
        data = json.loads(review_assignment.results.data)

        for relation in relations:
            field_name = core_logic.ascii_encode(relation.element.name)
            v = data[field_name]
            document.add_heading(relation.element.name, level=1)
            text = BeautifulSoup(
                (v[0]).encode('utf-8'),
                'html.parser'
            ).get_text()
            document.add_paragraph(text).bold = True
            recommendations = {
                'accept': 'Accept',
                'reject': 'Reject',
                'revisions': 'Revisions Required',
            }

        document.add_heading("Recommendation", level=1)
        if recommendations.get(review_assignment.recommendation):
            document.add_paragraph(
                recommendations[review_assignment.recommendation]
            ).italic = True
        document.add_heading("Competing Interests", level=1)
        document.add_paragraph(
            review_assignment.competing_interests
        ).italic = True

    else:
        p = document.add_paragraph(
            'You should complete this form and then use the review assignment '
            'page to upload it.'
        )

        for relation in relations:

            if (
                    relation.element.field_type in
                    ['text', 'textarea', 'date', 'email']
            ):
                document.add_heading(
                    relation.element.name + ": _______________________________",
                    level=1
                )
                document.add_paragraph(relation.help_text).italic = True

            elif relation.element.field_type in ['select', 'check']:
                document.add_heading(relation.element.name, level=1)

                if relation.element.field_type == 'select':
                    choices = render_choices(relation.element.choices)
                else:
                    choices = ['Y', 'N']

                p = document.add_paragraph(relation.help_text)
                p.add_run(
                    ' Mark your choice however you like, '
                    'as long as it is clear.'
                ).italic = True
                table = document.add_table(rows=2, cols=len(choices))
                hdr_cells = table.rows[0].cells

                for i, choice in enumerate(choices):
                    hdr_cells[i].text = choice[0]

                table.style = 'TableGrid'

    document.add_page_break()

    if not os.path.exists(os.path.join(settings.BASE_DIR, 'files', 'forms')):
        os.makedirs(os.path.join(settings.BASE_DIR, 'files', 'forms'))

    path = os.path.join(settings.FORM_DIR, f'{uuid4()}.docx')

    with default_storage.open(path, 'wb') as file_stream:
        document.save(file_stream)

    return path


def generate_review_form(
        request,
        review_type,
        submission_id,
        review_id,
        access_key=None,
):
    submission = get_object_or_404(core_models.Book, pk=submission_id)

    if access_key:
        review_assignment = get_object_or_404(
            core_models.ReviewAssignment,
            pk=review_id,
            access_key=access_key,
            review_type=review_type,
            withdrawn=False,
        )
    else:
        review_assignment = get_object_or_404(
            core_models.ReviewAssignment,
            pk=review_id,
            review_type=review_type,
            withdrawn=False,
        )

    path = create_completed_review_form(submission, review_assignment.pk)

    return serve_file(request, path)
