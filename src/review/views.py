import json
import mimetypes as mime
import os
from uuid import uuid4

from django.conf import settings
from django.contrib import messages
from django.contrib.auth.models import User
from django.core.urlresolvers import reverse
from django.db.models import Q
from django.http import Http404, StreamingHttpResponse, HttpResponseRedirect
from django.shortcuts import redirect, render, get_object_or_404
from django.utils import timezone
from django.utils.decorators import method_decorator
from django.utils.encoding import smart_text
from django.views.generic import FormView

from bs4 import BeautifulSoup
from docx import Document

from core import (
    email,
    forms as core_forms,
    log,
    logic as core_logic,
    models as core_models,
)
from core.decorators import is_reviewer, has_reviewer_role
from editorialreview import models as editorialreview_models
from core.files import handle_multiple_email_files
from review import (
    forms,
    logic,
    models,
)
from submission import models as submission_models
from core.setting_util import get_setting


@has_reviewer_role
def reviewer_dashboard(request):
    reopened_tasks = core_models.ReviewAssignment.objects.filter(
        user=request.user,
        completed__isnull=False,
        reopened=True,
        declined__isnull=True,
        withdrawn=False,
    ).select_related(
        'book',
    )
    incoming_tasks = core_models.ReviewAssignment.objects.filter(
        user=request.user,
        completed__isnull=True,
        reopened=False,
        declined__isnull=True,
        withdrawn=False
    ).select_related(
        'book',
    )

    pending_tasks = []

    [pending_tasks.append(task) for task in reopened_tasks]
    [pending_tasks.append(task) for task in incoming_tasks]

    completed_tasks = core_models.ReviewAssignment.objects.filter(
        user=request.user,
        completed__isnull=False,
        reopened=False,
        withdrawn=False
    ).select_related(
        'book',
    )
    pending_proposal_tasks = submission_models.ProposalReview.objects.filter(
        user=request.user,
        completed__isnull=True,
        declined__isnull=True,
        withdrawn=False,
    )
    completed_proposal_tasks = submission_models.ProposalReview.objects.filter(
        user=request.user,
        completed__isnull=False,
        withdrawn=False,
    )
    pending_submission_editorial_review_tasks = (
        editorialreview_models.EditorialReview.objects.filter(
            user=request.user,
            completed__isnull=True,
            withdrawn=False,
            content_type__model='book',
        )
    )
    pending_proposal_editorial_review_tasks = (
        editorialreview_models.EditorialReview.objects.filter(
            user=request.user,
            completed__isnull=True,
            withdrawn=False,
            content_type__model='proposal',
        )
    )
    completed_submission_editorial_review_tasks = (
        editorialreview_models.EditorialReview.objects.filter(
            user=request.user,
            completed__isnull=False,
            withdrawn=False,
            content_type__model='book',
        )
    )
    completed_proposal_editorial_review_tasks = (
        editorialreview_models.EditorialReview.objects.filter(
            user=request.user,
            completed__isnull=False,
            withdrawn=False,
            content_type__model='proposal',
        )
    )

    template = 'review/dashboard.html'
    context = {
        'pending_tasks': pending_tasks,
        'pending_proposal_tasks': pending_proposal_tasks,
        'pending_submission_editorial_review_tasks':
            pending_submission_editorial_review_tasks,
        'pending_proposal_editorial_review_tasks':
            pending_proposal_editorial_review_tasks,
        'pending_count': (
            len(pending_tasks)
            + len(pending_proposal_tasks)
            + len(pending_submission_editorial_review_tasks)
            + len(pending_proposal_editorial_review_tasks)
        ),
        'completed_tasks': completed_tasks,
        'completed_proposal_tasks': completed_proposal_tasks,
        'completed_submission_editorial_review_tasks':
            completed_submission_editorial_review_tasks,
        'completed_proposal_editorial_review_tasks':
            completed_proposal_editorial_review_tasks,
        'completed_count': (
            len(completed_tasks)
            + len(completed_proposal_tasks)
            + len(completed_submission_editorial_review_tasks)
            + len(completed_proposal_editorial_review_tasks)
        )
    }

    return render(request, template, context)


@is_reviewer
def reviewer_decision(
        request,
        review_type,
        submission_id,
        review_assignment_id,
        decision=None,
        access_key=None,
):
    """
    Check the review assignment has not been completed
    and is being accessed by the assigned user
    """
    submission = get_object_or_404(core_models.Book, pk=submission_id)
    one_click_no_login = get_setting('one_click_review_url', 'general')

    if one_click_no_login is not None:
        if one_click_no_login == 'on':
            one_click = True
            if access_key:
                review_assignment = get_object_or_404(
                    core_models.ReviewAssignment,
                    access_key=access_key,
                    pk=review_assignment_id,
                    declined__isnull=True,
                    review_type=review_type,
                    withdrawn=False,
                )
                user = review_assignment.user
            else:
                if request.user.is_authenticated():
                    review_assignment = get_object_or_404(
                        core_models.ReviewAssignment,
                        Q(user=request.user),
                        Q(book=submission),
                        Q(pk=review_assignment_id),
                        Q(declined__isnull=True),
                        Q(review_type=review_type),
                        Q(withdrawn=False),
                        Q(access_key__isnull=True) | Q(access_key__exact=''),
                    )
                    user = request.user
                else:
                    raise Http404
        else:
            one_click = False
            if access_key:
                review_assignment = get_object_or_404(
                    core_models.ReviewAssignment,
                    access_key=access_key,
                    pk=review_assignment_id,
                    declined__isnull=True,
                    review_type=review_type,
                    withdrawn=False,
                )
                user = review_assignment.user
            elif request.user.is_authenticated():
                review_assignment = get_object_or_404(
                    core_models.ReviewAssignment,
                    Q(user=request.user),
                    Q(book=submission),
                    Q(pk=review_assignment_id),
                    Q(declined__isnull=True),
                    Q(review_type=review_type),
                    Q(withdrawn=False),
                    Q(access_key__isnull=True) | Q(access_key__exact=''),
                )
                user = request.user
            else:
                raise Http404
    else:
        user = request.user
        if access_key:
            review_assignment = get_object_or_404(
                core_models.ReviewAssignment,
                access_key=access_key,
                pk=review_assignment_id,
                declined__isnull=True,
                review_type=review_type,
                withdrawn=False,
            )
        else:
            review_assignment = get_object_or_404(
                core_models.ReviewAssignment,
                Q(user=user),
                Q(book=submission),
                Q(pk=review_assignment_id),
                Q(declined__isnull=True),
                Q(review_type=review_type),
                Q(withdrawn=False),
                Q(access_key__isnull=True) | Q(access_key__exact=''),
            )

    book_editors = None
    series_editor = None
    if review_assignment:
        book_editors = review_assignment.book.book_editors.all()
        series_editor = review_assignment.book.get_series_editor()

    if review_assignment.accepted:
        if access_key:
            return redirect(reverse(
                'review_with_access_key',
                kwargs={
                    'review_type': review_type,
                    'submission_id': submission.pk,
                    'access_key': access_key,
                    'review_round': review_assignment.review_round.round_number
                }
            ))
        else:
            return redirect(reverse(
                'review_without_access_key',
                kwargs={
                    'review_type': review_type,
                    'submission_id': submission.pk,
                    'review_round': review_assignment.review_round.round_number
                }
            ))
    elif review_assignment.declined:
        return redirect(reverse('reviewer_dashboard'))

    # Form data POSTED to view takes precedence over 'decision' parameter
    if request.POST:
        if 'accept' in request.POST:
            decision = 'accept'
        elif 'decline' in request.POST:
            decision = 'decline'

    # Declaration of dictionaries containing redundant keyword arguments
    # used in calls to log function in review.logic
    editor_notification_kwargs = {
        'book': submission,
        'book_editors': book_editors,
        'series_editor': series_editor,
        'creator': user,
        'workflow': 'review'
    }
    decision_log_kwargs = {
        'book': submission,
        'user': user,
        'kind': 'review',
    }

    if decision == 'accept':
        review_assignment.accepted = timezone.now()
        message = (
            u"Review Assignment request for '{book_title}' has been "
            u"accepted by {first_name} {last_name}.".format(
                book_title=submission.title,
                first_name=review_assignment.user.first_name,
                last_name=review_assignment.user.last_name,
            )
        )

        decision_log_kwargs['message'] = message
        decision_log_kwargs['short_name'] = 'Assignment accepted'
        log.add_log_entry(**decision_log_kwargs)

        editor_notification_kwargs['message'] = message
        logic.notify_editors(**editor_notification_kwargs)

    elif decision == 'decline':
        review_assignment.declined = timezone.now()
        message = (
            u"Review Assignment request for '{book_title}' has been "
            u"declined by {first_name} {last_name}.".format(
                book_title=submission.title,
                first_name=review_assignment.user.first_name,
                last_name=review_assignment.user.last_name,
            )
        )

        decision_log_kwargs['message'] = message
        decision_log_kwargs['short_name'] = 'Assignment declined'
        log.add_log_entry(**decision_log_kwargs)

        editor_notification_kwargs['message'] = message
        logic.notify_editors(**editor_notification_kwargs)

    if decision:
        review_assignment.save()

        if access_key:
            return redirect(
                reverse(
                    'reviewer_decision_email_with_access_key',
                    kwargs={
                        'review_type': review_type,
                        'submission_id': submission.pk,
                        'review_assignment_id': review_assignment.pk,
                        'access_key': access_key,
                        'decision': decision,
                    }
                )
            )
        else:
            return redirect(
                reverse(
                    'reviewer_decision_email',
                    kwargs={
                        'review_type': review_type,
                        'submission_id': submission.pk,
                        'review_assignment_id': review_assignment.pk,
                        'decision': decision,
                    }
                )
            )

    template = 'review/reviewer_decision.html'
    context = {
        'submission': submission,
        'review_assignment': review_assignment,
        'has_additional_files': logic.has_additional_files(submission),
        'book_editors': book_editors,
        'series_editor': series_editor,
        'access_key': access_key,
        'one_click': one_click,
        'file_preview': get_setting('preview_review_files', 'general'),
        'instructions': get_setting('instructions_for_task_review', 'general')
    }

    return render(request, template, context)


class RequestedReviewerDecisionEmail(FormView):
    """Allows requested reviewers to send an email to requesting editors
    after they have responded to a peer review request.
    """

    template_name = 'shared/editable_notification_email.html'
    form_class = core_forms.CustomEmailForm

    @method_decorator(is_reviewer)
    def dispatch(self, request, *args, **kwargs):
        self.review_type = self.kwargs['review_type']
        self.review_assignment_id = self.kwargs['review_assignment_id']
        self.access_key = self.kwargs.get('access_key')
        self.decision = self.kwargs.get('decision')

        self.submission = get_object_or_404(
            core_models.Book,
            pk=self.kwargs['submission_id']
        )

        if self.access_key:
            self.review_assignment = get_object_or_404(
                core_models.ReviewAssignment,
                access_key=self.access_key,
                pk=self.review_assignment_id,
                review_type=self.review_type,
            )
        else:
            self.review_assignment = get_object_or_404(
                core_models.ReviewAssignment,
                Q(user=self.request.user),
                Q(book=self.submission),
                Q(pk=self.review_assignment_id),
                Q(review_type=self.review_type),
                Q(access_key__isnull=True) | Q(access_key__exact=''),
            )

        return super(RequestedReviewerDecisionEmail, self).dispatch(
            request,
            *args,
            **kwargs
        )

    def get_form_kwargs(self):
        """Renders the email body and subject for editing using the form."""
        kwargs = super(RequestedReviewerDecisionEmail, self).get_form_kwargs()

        if (
                self.review_assignment.assigning_editor and
                self.review_assignment.assigning_editor.profile.full_name()
        ):
            recipient_greeting = u'Dear {assigning_editor_name}'.format(
                assigning_editor_name=(
                    self.review_assignment.assigning_editor.profile.full_name()
                )
            )
        else:
            recipient_greeting = 'Dear sir or madam'

        email_context = {
            'greeting': recipient_greeting,
            'submission': self.submission,
            'sender': self.review_assignment.user,
        }

        if self.decision == 'accept':
            email_body = email.get_email_body(
                request=self.request,
                setting_name='requested_reviewer_accept',
                context=email_context,
            )
            email_subject = email.get_email_subject(
                request=self.request,
                setting_name='requested_reviewer_accept_subject',
                context=email_context,
            )
        else:
            email_body = email.get_email_body(
                request=self.request,
                setting_name='requested_reviewer_decline',
                context=email_context,
            )
            email_subject = email.get_email_subject(
                request=self.request,
                setting_name='requested_reviewer_decline_subject',
                context=email_context,
            )

        kwargs['initial'] = {
            'email_subject': email_subject,
            'email_body': email_body,
        }

        return kwargs

    def get_context_data(self, **kwargs):
        context = super(RequestedReviewerDecisionEmail, self).get_context_data(
            **kwargs
        )
        context['heading'] = (
            'Please ensure that you are happy with the below email to the '
            'editor notifying them of your decision'
        )
        return context

    def form_valid(self, form):
        attachments = handle_multiple_email_files(
            request_files=self.request.FILES.getlist('attachments'),
            file_owner=self.review_assignment.user,
        )
        assigning_editor_email = (
            self.review_assignment.assigning_editor.email
            if self.review_assignment.assigning_editor else ''
        )
        other_editors_emails = [
            editor.email
            for editor in self.submission.book_editors.exclude(
                email=assigning_editor_email,
                username=settings.INTERNAL_USER
            )
        ]

        series_editor = self.submission.get_series_editor()
        if series_editor and series_editor.email != assigning_editor_email:
            other_editors_emails.append(series_editor.email)

        from_email = self.review_assignment.user.email or get_setting(
            'from_address',
            'email'
        )
        email.send_prerendered_email(
            html_content=form.cleaned_data['email_body'],
            subject=form.cleaned_data['email_subject'],
            from_email=from_email,
            to=assigning_editor_email,
            cc=other_editors_emails,
            attachments=attachments,
            book=self.submission,
        )

        return super(RequestedReviewerDecisionEmail, self).form_valid(form)

    def get_success_url(self):
        if self.decision == 'accept':
            kwargs = {
                'review_type': self.review_type,
                'submission_id': self.submission.pk,
                'review_round':
                    self.review_assignment.review_round.round_number
            }
            if self.access_key:
                view_name = 'review_with_access_key'
                kwargs['access_key'] = self.access_key
            else:
                view_name = 'review_without_access_key'

            return reverse(view_name, kwargs=kwargs)

        return reverse('review_request_declined')


@is_reviewer
def review(request, review_type, submission_id, review_round, access_key=None):
    ci_required = get_setting('ci_required', 'general')
    one_click_no_login = core_models.Setting.objects.filter(
        name='one_click_review_url',
    )

    if one_click_no_login[0].value == 'on' and access_key:
        one_click = True
        review_assignment = get_object_or_404(
            core_models.ReviewAssignment,
            access_key=access_key,
            review_round__round_number=review_round,
            declined__isnull=True,
            review_type=review_type,
            withdrawn=False,
        )
        submission = get_object_or_404(core_models.Book, pk=submission_id)
        user = review_assignment.user

        if review_assignment.completed and not review_assignment.reopened:
            return redirect(reverse(
                'review_complete_with_access_key',
                kwargs={
                    'review_type': review_type,
                    'submission_id': submission.pk,
                    'access_key': access_key,
                    'review_round': review_round
                }
            ))

    else:
        one_click = False
        user = request.user

        if access_key:
            review_assignment = get_object_or_404(
                core_models.ReviewAssignment,
                access_key=access_key,
                review_round__round_number=review_round,
                declined__isnull=True,
                review_type=review_type,
                withdrawn=False,
            )
            submission = get_object_or_404(core_models.Book, pk=submission_id)

            if review_assignment.completed and not review_assignment.reopened:
                return redirect(reverse(
                    'review_complete_with_access_key',
                    kwargs={
                        'review_type': review_type,
                        'submission_id': submission.pk,
                        'access_key': access_key,
                        'review_round': review_round
                    }
                ))

        elif review_type == 'proposal':
            submission = get_object_or_404(
                submission_models.Proposal,
                pk=submission_id,
            )
            review_assignment = get_object_or_404(
                submission_models.ProposalReview,
                user=user,
                proposal=submission,
                completed__isnull=True,
                declined__isnull=True,
                withdrawn=False,
            )

        else:
            submission = get_object_or_404(core_models.Book, pk=submission_id)
            review_assignment = get_object_or_404(
                core_models.ReviewAssignment,
                Q(user=user),
                Q(book=submission),
                Q(review_round__round_number=review_round),
                Q(declined__isnull=True),
                Q(review_type=review_type),
                Q(withdrawn=False),
                Q(access_key__isnull=True) | Q(access_key__exact=''),
            )

            if review_assignment.completed and not review_assignment.reopened:
                return redirect(reverse(
                    'review_complete',
                    kwargs={
                        'review_type': review_type,
                        'submission_id': submission.pk,
                        'review_round': review_round
                    }
                ))

    book_editors = review_assignment.book.book_editors.all()
    series_editor = review_assignment.book.get_series_editor()

    # Check that this review is being access by the user,
    # is not completed and has not been declined.
    if review_assignment:
        if not review_assignment.accepted and not review_assignment.declined:
            if access_key:
                return redirect(reverse(
                    'reviewer_decision_without_access_key',
                    kwargs={
                         'review_type': review_type,
                         'submission_id': submission.pk,
                         'access_key': access_key,
                         'review_assignment_id': review_assignment.pk
                    }
                ))
            else:
                return redirect(reverse(
                    'reviewer_decision_without',
                    kwargs={
                        'review_type': review_type,
                        'submission_id': submission.pk,
                        'review_assignment_id': review_assignment.pk
                    }
                ))

    if review_assignment.review_form:
        form = forms.GeneratedForm(form=review_assignment.review_form)
    else:
        review_assignment.review_form = submission.review_form
        review_assignment.save()
        form = forms.GeneratedForm(form=submission.review_form)

    if review_assignment.reopened:
        result = review_assignment.results
        if result:
            initial_data = {}
            data = json.loads(result.data)
            for k, v in data.items():
                initial_data[k] = v[0]

            form.initial = initial_data

    recommendation_form = core_forms.RecommendationForm(
        ci_required=ci_required
    )

    if review_assignment.reopened:
        initial_data = {
            u'recommendation': review_assignment.recommendation,
            u'competing_interests': review_assignment.competing_interests,
        }
        recommendation_form.initial = initial_data

    if not request.POST and request.GET.get('download') == 'docx':
        path = create_completed_review_form(submission, review_assignment.pk)
        return serve_file(request, path)
    elif request.POST:
        form = forms.GeneratedForm(
            request.POST,
            request.FILES,
            form=review_assignment.review_form,
        )
        recommendation_form = core_forms.RecommendationForm(
            request.POST,
            ci_required=ci_required,
        )
        if form.is_valid() and recommendation_form.is_valid():
            save_dict = {}
            file_fields = models.FormElementsRelationship.objects.filter(
                form=review_assignment.review_form,
                element__field_type='upload',
            )
            data_fields = models.FormElementsRelationship.objects.filter(
                ~Q(element__field_type='upload'),
                form=review_assignment.review_form,
            )

            for field in file_fields:
                field_name = core_logic.ascii_encode(field.element.name)
                if field_name in request.FILES:
                    save_dict[field_name] = [
                        logic.handle_review_file(
                            request.FILES[field_name],
                            'book',
                            review_assignment,
                            'reviewer'
                        )
                    ]

            for field in data_fields:
                field_name = core_logic.ascii_encode(field.element.name)
                if field_name in request.POST:
                    save_dict[field_name] = [
                        request.POST.get(field_name),
                        'text',
                    ]

            json_data = smart_text(json.dumps(save_dict))

            if review_assignment.reopened:
                if review_assignment.results:
                    review_assignment.results.data = json_data
                    review_assignment.results.save()
                    review_assignment.reopened = False
                    review_assignment.save()
                else:
                    form_results = models.FormResult(
                        form=review_assignment.review_form,
                        data=json_data,
                    )
                    form_results.save()
                    review_assignment.results = form_results
                    review_assignment.reopened = False
                    review_assignment.save()
            else:
                form_results = models.FormResult(
                    form=review_assignment.review_form,
                    data=json_data,
                )
                form_results.save()
                review_assignment.results = form_results
                review_assignment.save()

            if request.FILES.get('review_file_upload'):
                logic.handle_review_file(
                    request.FILES.get('review_file_upload'),
                    'book',
                    review_assignment,
                    'reviewer'
                )

            review_assignment.completed = timezone.now()
            if not review_assignment.accepted:
                review_assignment.accepted = timezone.now()
            review_assignment.recommendation = request.POST.get(
                'recommendation',
            )
            review_assignment.competing_interests = request.POST.get(
                'competing_interests',
            )
            review_assignment.save()

            message = (
                "%s Review assignment with id %s has been completed by %s ." % (
                    review_assignment.review_type.title(),
                    review_assignment.id,
                    review_assignment.user.profile.full_name()
                )
            )
            press_editors = User.objects.filter(
                profile__roles__slug='press-editor',
            )

            for editor in press_editors:
                notification = core_models.Task(
                    assignee=editor,
                    creator=user,
                    text=message,
                    workflow='review',
                    book=submission
                )
                notification.save()

            if not review_type == 'proposal':
                log.add_log_entry(
                    book=submission,
                    user=user,
                    kind='review',
                    message='Reviewer %s %s completed review for %s.' % (
                        review_assignment.user.first_name,
                        review_assignment.user.last_name,
                        submission.title
                    ),
                    short_name='Assignment Completed',
                )
                message = "Reviewer %s %s has completed a review for '%s'." % (
                    submission.title,
                    review_assignment.user.first_name,
                    review_assignment.user.last_name,
                )
                logic.notify_editors(
                    book=submission,
                    message=message,
                    book_editors=book_editors,
                    series_editor=series_editor,
                    creator=user,
                    workflow='review'
                )

            if access_key:
                return redirect(
                    reverse(
                        'review_completion_email_with_access_key',
                        kwargs={
                            'review_type': review_type,
                            'submission_id': submission.id,
                            'review_assignment_id': review_assignment.id,
                            'access_key': access_key,
                            'review_round': review_round
                        }
                    )
                )
            else:
                return redirect(
                    reverse(
                        'review_completion_email',
                        kwargs={
                            'review_type': review_type,
                            'submission_id': submission.id,
                            'review_assignment_id': review_assignment.id,
                            'review_round': review_round
                        }
                    )
                )

    template = 'review/review.html'
    context = {
        'review_assignment': review_assignment,
        'submission': submission,
        'form': form,
        'form_info': review_assignment.review_form,
        'recommendation_form': recommendation_form,
        'book_editors': book_editors,
        'series_editor': series_editor,
        'access_key': access_key,
        'one_click': one_click,
        'has_additional_files': logic.has_additional_files(submission),
        'instructions': get_setting('instructions_for_task_review', 'general')
    }

    return render(request, template, context)


def review_request_declined(request):
    template = 'review/review_request_declined.html'
    return render(request, template)


class ReviewCompletionEmail(FormView):
    """Allows peer reviewers who have just completed a review to
    customise a notification email before it is sent to the editors.
    """

    template_name = 'shared/editable_notification_email.html'
    form_class = core_forms.CustomEmailForm

    @method_decorator(is_reviewer)
    def dispatch(self, request, *args, **kwargs):
        self.review_type = self.kwargs['review_type']
        self.submission_id = self.kwargs['submission_id']
        self.review_round = self.kwargs['review_round']
        self.access_key = self.kwargs.get('access_key')

        self.submission = get_object_or_404(
            core_models.Book,
            pk=self.submission_id
        )
        self.review_assignment = get_object_or_404(
            core_models.ReviewAssignment,
            pk=self.kwargs.get('review_assignment_id')
        )

        self.recipient_editors = []
        self.recipient_editors.extend(self.submission.book_editors.all())
        series_editor = self.submission.get_series_editor()
        if series_editor:
            self.recipient_editors.append(series_editor)

        return super(ReviewCompletionEmail, self).dispatch(
            request,
            *args,
            **kwargs
        )

    def get_form_kwargs(self):
        """Renders the email body and subject for editing using the form"""
        kwargs = super(ReviewCompletionEmail, self).get_form_kwargs()

        if self.review_assignment.assigning_editor:
            recipient_greeting = u'Dear {name}'.format(
                name=self.review_assignment.assigning_editor.profile.full_name()
            )
        else:
            recipient_greeting = 'Dear sir or madam'

        email_context = {
            'greeting': recipient_greeting,
            'submission': self.submission,
            'sender': self.review_assignment.user,
        }

        kwargs['initial'] = {
            'email_subject': email.get_email_subject(
                request=self.request,
                setting_name='peer_review_completed_subject',
                context=email_context,
            ),
            'email_body': email.get_email_body(
                request=self.request,
                setting_name='peer_review_completed',
                context=email_context,
            ),
        }

        return kwargs

    def get_context_data(self, **kwargs):
        context = super(ReviewCompletionEmail, self).get_context_data(**kwargs)
        context['heading'] = (
            'Please ensure that you are happy with the below email to the '
            'editor notifying them that you have completed your review'
        )
        return context

    def form_valid(self, form):
        attachments = handle_multiple_email_files(
            request_files=self.request.FILES.getlist('attachments'),
            file_owner=self.review_assignment.user,
        )

        if self.review_assignment.assigning_editor:
            assigning_editor_email = (
                self.review_assignment.assigning_editor.email
            )
        else:
            assigning_editor_email = ''

        other_editors_emails = [
            editor.email
            for editor in self.submission.book_editors.exclude(
                email=assigning_editor_email,
                username=settings.INTERNAL_USER
            )
        ]
        series_editor = self.submission.get_series_editor()
        if series_editor and series_editor.email != assigning_editor_email:
            other_editors_emails.append(series_editor.email)

        from_email = self.review_assignment.user.email or get_setting(
            'from_address',
            'email'
        )
        email.send_prerendered_email(
            html_content=form.cleaned_data['email_body'],
            subject=form.cleaned_data['email_subject'],
            from_email=from_email,
            to=assigning_editor_email,
            cc=other_editors_emails,
            attachments=attachments,
            book=self.submission,
        )

        return super(ReviewCompletionEmail, self).form_valid(form)

    def get_success_url(self):
        if self.access_key:
            return reverse(
                'review_complete_with_access_key',
                kwargs={
                    'review_type': self.review_type,
                    'submission_id': self.submission.pk,
                    'access_key': self.access_key,
                    'review_round': self.review_round,
                }
            )

        return reverse(
            'review_complete',
            kwargs={
                'review_type': self.review_type,
                'submission_id': self.submission.pk,
                'review_round': self.review_round,
            }
        )


@is_reviewer
def review_complete(
        request,
        review_type,
        submission_id,
        review_round,
        access_key=None,
):
    one_click_no_login = core_models.Setting.objects.filter(
        name='one_click_review_url',
    )

    if one_click_no_login[0].value == 'on' and access_key:
        review_assignment = get_object_or_404(
            core_models.ReviewAssignment,
            access_key=access_key,
        )
        submission = get_object_or_404(core_models.Book, pk=submission_id)
        user = review_assignment.user
    else:
        user = request.user

        if access_key:
            review_assignment = get_object_or_404(
                core_models.ReviewAssignment,
                access_key=access_key,
                declined__isnull=True,
                review_type=review_type,
                review_round=review_round,
                withdrawn=False,
            )
            submission = get_object_or_404(core_models.Book, pk=submission_id)

        elif review_type == 'proposal':
            submission = get_object_or_404(
                submission_models.Proposal,
                pk=submission_id,
            )
            review_assignment = get_object_or_404(
                submission_models.ProposalReview,
                user=user,
                proposal=submission,
            )

        else:
            submission = get_object_or_404(core_models.Book, pk=submission_id)
            review_assignment = get_object_or_404(
                core_models.ReviewAssignment,
                Q(user=user),
                Q(review_round__round_number=review_round),
                Q(book=submission),
                Q(withdrawn=False),
                Q(review_type=review_type),
                Q(access_key__isnull=True) | Q(access_key__exact=''),
            )

    result = review_assignment.results

    if not result or not review_assignment.completed:
        if not result:
            review_assignment.completed = None
            review_assignment.save()
        if access_key:
            return redirect(reverse(
                'review_with_access_key',
                kwargs={
                    'review_type': review_type,
                    'submission_id': submission.id,
                    'access_key': access_key,
                    'review_round': review_round
                }
            ))
        else:
            return redirect(reverse(
                'review_without_access_key',
                kwargs={
                    'review_type': review_type,
                    'submission_id': submission.id,
                    'review_round': review_round
                }
            ))

    elif review_assignment.completed and review_assignment.reopened:
        if access_key:
            return redirect(reverse(
                'review_with_access_key',
                kwargs={
                    'review_type': review_type,
                    'submission_id': submission.id,
                    'access_key': access_key,
                    'review_round': review_round
                }
            ))
        else:
            return redirect(reverse(
                'review_without_access_key',
                kwargs={
                    'review_type': review_type,
                    'submission_id': submission.id,
                    'review_round': review_round
                }
            ))

    if not request.POST and request.GET.get('download') == 'docx':
        path = create_completed_review_form(submission, review_assignment.pk)
        return serve_file(request, path)

    relations = models.FormElementsRelationship.objects.filter(form=result.form)
    data_ordered = core_logic.order_data(
        core_logic.decode_json(result.data),
        relations,
    )

    if not review_assignment.review_form:
        review_assignment.review_form = submission.review_form
        review_assignment.save()

    template = 'review/complete.html'
    context = {
        'submission': submission,
        'review_assignment': review_assignment,
        'form_info': review_assignment.review_form,
        'data_ordered': data_ordered,
        'result': result,
        'additional_files': logic.has_additional_files(submission),
        'book_editors': review_assignment.book.book_editors.all(),
        'series_editor': review_assignment.book.get_series_editor(),
        'instructions': get_setting('instructions_for_task_review', 'general')
    }

    return render(request, template, context)


@is_reviewer
def review_complete_no_redirect(
        request,
        review_type,
        submission_id,
        review_round,
        access_key=None,
):
    user = request.user
    if access_key:
        review_assignment = get_object_or_404(
            core_models.ReviewAssignment,
            access_key=access_key,
            results__isnull=False,
            declined__isnull=True,
            review_type=review_type,
            review_round=review_round,
            withdrawn=False,
        )
        submission = get_object_or_404(core_models.Book, pk=submission_id)
    elif review_type == 'proposal':
        submission = get_object_or_404(
            submission_models.Proposal,
            pk=submission_id,
        )
        review_assignment = get_object_or_404(
            submission_models.ProposalReview,
            user=user,
            proposal=submission,
            results__isnull=False,
        )
    else:
        submission = get_object_or_404(core_models.Book, pk=submission_id)
        review_assignment = get_object_or_404(
            core_models.ReviewAssignment,
            Q(user=user),
            Q(results__isnull=False),
            Q(review_round__round_number=review_round),
            Q(book=submission),
            Q(withdrawn=False),
            Q(review_type=review_type),
            Q(access_key__isnull=True) | Q(access_key__exact='')
        )

    result = review_assignment.results

    if not request.POST and request.GET.get('download') == 'docx':
        path = create_completed_review_form(submission, review_assignment.pk)
        return serve_file(request, path)

    relations = models.FormElementsRelationship.objects.filter(form=result.form)
    data_ordered = core_logic.order_data(
        core_logic.decode_json(result.data),
        relations,
    )

    if not review_assignment.review_form:
        review_assignment.review_form = submission.review_form
        review_assignment.save()

    template = 'review/complete.html'
    context = {
        'submission': submission,
        'review_assignment': review_assignment,
        'form_info': review_assignment.review_form,
        'data_ordered': data_ordered,
        'result': result,
        'additional_files': logic.has_additional_files(submission),
        'book_editors': review_assignment.book.book_editors.all(),
        'series_editor': review_assignment.book.get_series_editor(),
        'instructions': get_setting('instructions_for_task_review', 'general')
    }

    return render(request, template, context)


def editorial_review(request, submission_id, access_key):
    ci_required = get_setting('ci_required', 'general')
    # Check user has accessed review, is not completed and not declined.
    review_assignment = get_object_or_404(
        core_models.EditorialReviewAssignment,
        Q(publishing_committee_access_key=access_key) |
        Q(editorial_board_access_key=access_key),
    )
    submission = get_object_or_404(core_models.Book, pk=submission_id)

    resubmit = True

    if review_assignment.completed and not review_assignment.reopened:
        return redirect(reverse(
            'editorial_review_complete',
            kwargs={'submission_id': submission.pk, 'access_key': access_key}
        ))
    editorial_board = False

    if access_key == review_assignment.editorial_board_access_key:
        editorial_board = True
    elif access_key == review_assignment.publishing_committee_access_key:
        editorial_board = False

    if editorial_board:
        editorial_result = None
        editorial_data_ordered = None
        form_info = review_assignment.editorial_board_review_form
        form = forms.GeneratedForm(
            form=review_assignment.editorial_board_review_form,
        )
    else:
        editorial_result = review_assignment.editorial_board_results
        editorial_relations = models.FormElementsRelationship.objects.filter(
            form=editorial_result.form,
        )
        editorial_data_ordered = core_logic.order_data(
            core_logic.decode_json(editorial_result.data),
            editorial_relations,
        )

        if not review_assignment.publication_committee_review_form:
            form_info = review_assignment.editorial_board_review_form
            form = forms.GeneratedForm(
                form=review_assignment.editorial_board_review_form,
            )
        else:
            form_info = review_assignment.publication_committee_review_form
            form = forms.GeneratedForm(
                form=review_assignment.publication_committee_review_form,
            )

    if editorial_board:
        result = review_assignment.editorial_board_results
        if result and not review_assignment.editorial_board_passed:
            initial_data = {}
            data = json.loads(result.data)
            for k, v in data.items():
                initial_data[k] = v[0]
            form.initial = initial_data
        elif not result:
            resubmit = True
        elif result and review_assignment.editorial_board_passed:
            resubmit = False
    else:
        result = review_assignment.publication_committee_results
        if result and not review_assignment.publication_committee_passed:
            initial_data = {}
            data = json.loads(result.data)
            for k, v in data.items():
                initial_data[k] = v[0]
            form.initial = initial_data
        elif not result:
            resubmit = True
        elif result and review_assignment.publication_committee_passed:
            resubmit = False

    recommendation_form = core_forms.RecommendationForm(
        ci_required=ci_required,
    )

    initial_data = {}
    if editorial_board and not review_assignment.editorial_board_passed:
        initial_data[u'recommendation'] = (
            review_assignment.editorial_board_recommendation
        )
        initial_data[
            u'competing_interests'] = (
            review_assignment.editorial_board_competing_interests
        )
    elif (
        not editorial_board and
        not review_assignment.publication_committee_passed
    ):
        initial_data[
            u'recommendation'] = (
            review_assignment.publication_committee_recommendation
        )
        initial_data[
            u'competing_interests'] = (
            review_assignment.publication_committee_competing_interests
        )

    recommendation_form.initial = initial_data

    if not request.POST and request.GET.get('download') == 'docx':
        if editorial_board:
            path = create_review_form(
                submission,
                review_assignment.editorial_board_review_form,
            )
        else:
            path = create_review_form(
                submission,
                review_assignment.publication_committee_review_form,
            )

        return serve_file(request, path)
    elif request.POST:
        if editorial_board:
            form = forms.GeneratedForm(
                request.POST,
                request.FILES,
                form=review_assignment.editorial_board_review_form,
            )
        else:
            if not review_assignment.publication_committee_review_form:
                form = forms.GeneratedForm(
                    request.POST,
                    request.FILES,
                    form=review_assignment.editorial_board_review_form,
                )
            else:
                form = forms.GeneratedForm(
                    request.POST,
                    request.FILES,
                    form=review_assignment.publication_committee_review_form,
                )

        recommendation_form = core_forms.RecommendationForm(
            request.POST,
            ci_required=ci_required,
        )

        if form.is_valid() and recommendation_form.is_valid():
            save_dict = {}
            if editorial_board:
                file_fields = models.FormElementsRelationship.objects.filter(
                    form=review_assignment.editorial_board_review_form,
                    element__field_type='upload',
                )
                data_fields = models.FormElementsRelationship.objects.filter(
                    ~Q(element__field_type='upload'),
                    form=review_assignment.editorial_board_review_form,
                )
            else:
                if not review_assignment.publication_committee_review_form:
                    file_fields = (
                        models.FormElementsRelationship.objects.filter(
                            form=review_assignment.editorial_board_review_form,
                            element__field_type='upload',
                        )
                    )
                    data_fields = (
                        models.FormElementsRelationship.objects.filter(
                            ~Q(element__field_type='upload'),
                            form=review_assignment.editorial_board_review_form,
                        )
                    )
                else:
                    file_fields = (
                        models.FormElementsRelationship.objects.filter(
                            form=review_assignment.
                            publication_committee_review_form,
                            element__field_type='upload',
                        )
                    )
                    data_fields = (
                        models.FormElementsRelationship.objects.filter(
                            ~Q(element__field_type='upload'),
                            form=review_assignment.
                            publication_committee_review_form,
                        )
                    )

            for field in file_fields:
                if field.element.name in request.FILES:
                    # TODO change value from string to list [value, value_type].
                    save_dict[field.element.name] = [
                        handle_editorial_review_file(
                            request.FILES[field.element.name], submission,
                            review_assignment, 'reviewer', editorial_board
                        )
                    ]

            for field in data_fields:
                if field.element.name in request.POST:
                    # TODO change value from string to list [value, value_type].
                    save_dict[field.element.name] = [
                        request.POST.get(field.element.name),
                        'text',
                    ]

            data = smart_text(json.dumps(save_dict))
            if review_assignment.reopened:

                if editorial_board:
                    review_assignment.editorial_board_results.data = data
                    review_assignment.editorial_board_results.save()
                    review_assignment.reopened = False
                    review_assignment.save()
                else:
                    review_assignment.publication_committee_results.data = data
                    review_assignment.publication_committee_results.save()
                    review_assignment.reopened = False
                    review_assignment.save()

            else:
                if editorial_board:
                    result = models.FormResult(
                        form=review_assignment.editorial_board_review_form,
                        data=data,
                    )
                    result.save()
                    review_assignment.editorial_board_results = result
                    review_assignment.save()
                else:
                    if not review_assignment.publication_committee_review_form:
                        result = models.FormResult(
                            form=review_assignment.editorial_board_review_form,
                            data=data,
                        )
                    else:
                        result = models.FormResult(
                            form=review_assignment.
                            publication_committee_review_form,
                            data=data,
                        )
                    result.save()
                    review_assignment.publication_committee_results = result
                    review_assignment.save()

            if request.FILES.get('review_file_upload'):
                handle_editorial_review_file(
                    request.FILES.get('review_file_upload'), submission,
                    review_assignment, 'reviewer', editorial_board)

            if editorial_board:
                review_assignment.editorial_board_recommendation = (
                    request.POST.get('recommendation')
                )
                review_assignment.editorial_board_competing_interests = (
                    request.POST.get('competing_interests')
                )
            else:
                review_assignment.publication_committee_recommendation = (
                    request.POST.get('recommendation')
                )
                review_assignment.publication_committee_competing_interests = (
                    request.POST.get('competing_interests')
                )

            review_assignment.save()
            message = "Editorial review assignment #%s has been completed." % (
                review_assignment.id,
            )
            press_editors = User.objects.filter(
                profile__roles__slug='press-editor',
            )

            for editor in press_editors:
                notification = core_models.Task(
                    assignee=editor,
                    creator=request.user,
                    text=message,
                    workflow='editorial-review',
                    book=submission,
                    editorial_review=review_assignment,
                )
                notification.save()

            messages.add_message(
                request,
                messages.INFO,
                'Submitted successfully',
            )

            if editorial_board:
                return redirect(reverse(
                    'editorial_review',
                    kwargs={
                        'submission_id': submission.id,
                        'access_key':
                        review_assignment.editorial_board_access_key
                    }
                ))
            else:
                return redirect(reverse(
                    'editorial_review',
                    kwargs={
                        'submission_id': submission.id,
                        'access_key':
                        review_assignment.publishing_committee_access_key
                    }
                ))

    template = 'review/editorial_review.html'
    context = {
        'review_assignment': review_assignment,
        'submission': submission,
        'form': form,
        'editorial_board': editorial_board,
        'form_info': form_info,
        'recommendation_form': recommendation_form,
        'book_editors': review_assignment.book.book_editors.all(),
        'series_editor': review_assignment.book.get_series_editor(),
        'resubmit': resubmit,
        'editorial_data_ordered': editorial_data_ordered,
        'editorial_result': editorial_result,
        'has_additional_files': logic.has_additional_files(submission),
        'instructions': get_setting('instructions_for_task_review', 'general')
    }

    return render(request, template, context)


def editorial_review_complete(request, submission_id, access_key):
    review_assignment = get_object_or_404(
        core_models.EditorialReviewAssignment,
        Q(publishing_committee_access_key=access_key) |
        Q(editorial_board_access_key=access_key)
    )
    submission = get_object_or_404(core_models.Book, pk=submission_id)

    if access_key == review_assignment.editorial_board_access_key:
        editorial_board = True
        result = review_assignment.editorial_board_results
        review_form = review_assignment.editorial_board_review_form
    elif access_key == review_assignment.publishing_committee_access_key:
        editorial_board = False
        result = review_assignment.publication_committee_results
        review_form = review_assignment.publication_committee_review_form

    if not result and not editorial_board:
        return redirect(reverse(
            'editorial_review',
            kwargs={
                'submission_id': submission.id,
                'access_key': review_assignment.publishing_committee_access_key
            }
        ))
    elif not result and editorial_board:
        return redirect(reverse(
            'editorial_review',
            kwargs={
                'submission_id': submission.id,
                'access_key': review_assignment.editorial_board_access_key
            }
        ))

    elif review_assignment.completed and review_assignment.reopened:
        if editorial_board:
            return redirect(reverse(
                'editorial_review_complete',
                kwargs={
                    'submission_id': submission.id,
                    'access_key':
                    review_assignment.publishing_committee_access_key
                }
            ))
        else:
            return redirect(reverse(
                'editorial_review_complete',
                kwargs={
                    'submission_id': submission.id,
                    'access_key': review_assignment.editorial_board_access_key
                }
            ))

    relations = models.FormElementsRelationship.objects.filter(form=result.form)
    data_ordered = core_logic.order_data(
        core_logic.decode_json(result.data),
        relations
    )

    template = 'review/editorial_complete.html'
    context = {
        'submission': submission,
        'review_assignment': review_assignment,
        'form_info': review_form,
        'data_ordered': data_ordered,
        'result': result,
        'editorial_board': editorial_board,
        'additional_files': logic.has_additional_files(submission),
        'book_editors': review_assignment.book.book_editors.all(),
        'series_editor': review_assignment.book.get_series_editor(),
        'instructions': get_setting('instructions_for_task_review', 'general')
    }

    return render(request, template, context)


def render_choices(choices):
    c_split = choices.split('|')
    return [(choice.capitalize(), choice) for choice in c_split]


def create_completed_review_form(submission, review_id):
    document = Document()
    document.add_heading(submission.title, 0)
    review_assignment = get_object_or_404(
        core_models.ReviewAssignment,
        pk=review_id,
    )
    if review_assignment.review_form:
        relations = models.FormElementsRelationship.objects.filter(
            form=review_assignment.review_form,
        ).order_by('order')
    else:
        review_assignment.review_form = submission.review_form
        review_assignment.save()
        relations = models.FormElementsRelationship.objects.filter(
            form=submission.review_form,
        ).order_by('order')

    if review_assignment.results:
        data = json.loads(review_assignment.results.data)

        for relation in relations:
            field_name = core_logic.ascii_encode(relation.element.name)
            v = data[field_name]
            document.add_heading(relation.element.name, level=1)
            text = BeautifulSoup(
                (v[0]).encode('utf-8'),
                'html.parser'
            ).get_text()
            document.add_paragraph(text).bold = True
            recommendations = {
                'accept': 'Accept',
                'reject': 'Reject',
                'revisions': 'Revisions Required',
            }

        document.add_heading("Recommendation", level=1)
        if recommendations.get(review_assignment.recommendation):
            document.add_paragraph(
                recommendations[review_assignment.recommendation]
            ).italic = True
        document.add_heading("Competing Interests", level=1)
        document.add_paragraph(
            review_assignment.competing_interests
        ).italic = True

    else:
        p = document.add_paragraph(
            'You should complete this form and then use the review assignment '
            'page to upload it.'
        )

        for relation in relations:

            if (
                relation.element.field_type
                in ['text', 'textarea', 'date', 'email']
            ):
                document.add_heading(
                    relation.element.name + ": _______________________________",
                    level=1
                )
                document.add_paragraph(relation.help_text).italic = True

            if relation.element.field_type in ['select', 'check']:
                document.add_heading(relation.element.name, level=1)

                if relation.element.field_type == 'select':
                    choices = render_choices(relation.element.choices)
                else:
                    choices = ['Y', 'N']

                p = document.add_paragraph(relation.help_text)
                p.add_run(
                    ' Mark your choice however you like, '
                    'as long as it is clear.'
                ).italic = True
                table = document.add_table(rows=2, cols=len(choices))
                hdr_cells = table.rows[0].cells

                for i, choice in enumerate(choices):
                    hdr_cells[i].text = choice[0]

                table.style = 'TableGrid'

    document.add_page_break()

    if not os.path.exists(os.path.join(settings.BASE_DIR, 'files', 'forms')):
        os.makedirs(os.path.join(settings.BASE_DIR, 'files', 'forms'))

    path = os.path.join(
        settings.BASE_DIR,
        'files',
        'forms',
        '%s.docx' % str(uuid4())
    )
    document.save(path)

    return path


def generate_review_form(
        request,
        review_type,
        submission_id,
        review_id,
        access_key=None,
):
    submission = get_object_or_404(core_models.Book, pk=submission_id)

    if access_key:
        review_assignment = get_object_or_404(
            core_models.ReviewAssignment,
            pk=review_id,
            access_key=access_key,
            review_type=review_type,
            withdrawn=False,
        )
    else:
        review_assignment = get_object_or_404(
            core_models.ReviewAssignment,
            pk=review_id,
            review_type=review_type,
            withdrawn=False,
        )

    path = create_completed_review_form(submission, review_assignment.pk)

    return serve_file(request, path)


def create_review_form(submission, review_form):
    document = Document()
    document.add_heading(submission.title, 0)
    p = document.add_paragraph(
        'You should complete this form and then use the review '
        'page to upload it.'
    )
    relations = models.FormElementsRelationship.objects.filter(
        form=review_form
    ).order_by('order')

    for relation in relations:
        if relation.element.field_type in ['text', 'textarea', 'date', 'email']:
            document.add_heading(
                relation.element.name + ": _______________________________",
                level=1,
            )
            document.add_paragraph(relation.help_text).italic = True

        if relation.element.field_type in ['select', 'check']:
            document.add_heading(relation.element.name, level=1)

            if relation.element.field_type == 'select':
                choices = render_choices(relation.element.choices)
            else:
                choices = ['Y', 'N']

            p = document.add_paragraph(relation.help_text)
            p.add_run(
                ' Mark your choice however you like, as long as it is clear.'
            ).italic = True
            table = document.add_table(rows=2, cols=len(choices))
            hdr_cells = table.rows[0].cells

            for i, choice in enumerate(choices):
                hdr_cells[i].text = choice[0]

            table.style = 'TableGrid'

    document.add_page_break()

    if not os.path.exists(os.path.join(settings.BASE_DIR, 'files', 'forms')):
        os.makedirs(os.path.join(settings.BASE_DIR, 'files', 'forms'))

    path = os.path.join(
        settings.BASE_DIR,
        'files',
        'forms',
        '%s.docx' % str(uuid4())
    )

    document.save(path)

    return path


@is_reviewer
def serve_file(request, file_path):
    try:
        fsock = open(file_path, 'r')
        mimetype = mime.guess_type(file_path)
        response = StreamingHttpResponse(fsock, content_type=mimetype)
        response['Content-Disposition'] = (
            "attachment; filename=review_form.docx"
        )
        return response
    except IOError:
        messages.add_message(request, messages.ERROR, 'File not found.')
        return HttpResponseRedirect(request.META.get('HTTP_REFERER'))


def handle_editorial_review_file(
        file,
        proposal,
        review_assignment,
        kind,
        editorial
):
    original_filename = smart_text(file._get_name())
    filename = str(uuid4()) + str(os.path.splitext(original_filename)[1])
    folder_structure = os.path.join(
        settings.BASE_DIR,
        'files',
        'books',
        str(review_assignment.book.id)
    )

    if not os.path.exists(folder_structure):
        os.makedirs(folder_structure)

    path = os.path.join(folder_structure, str(filename))
    fd = open(path, 'wb')

    for chunk in file.chunks():
        fd.write(chunk)

    fd.close()
    file_mime = mime.guess_type(filename)

    try:
        file_mime = file_mime[0]
        if not file_mime:
            file_mime = 'unknown'
    except IndexError:
        file_mime = 'unknown'

    new_file = core_models.File(
        mime_type=file_mime,
        original_filename=original_filename,
        uuid_filename=filename,
        stage_uploaded=1,
        kind=kind,
        owner=review_assignment.management_editor,
    )
    new_file.save()

    if editorial:
        review_assignment.editorial_board_files.add(new_file)
    else:
        review_assignment.publication_committee_files.add(new_file)

    return path
