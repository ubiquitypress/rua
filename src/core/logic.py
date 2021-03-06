import os
from builtins import any as string_any
import datetime
import json
import mimetypes
import re
from uuid import uuid4

from bs4 import BeautifulSoup
from django.conf import settings
from django.contrib import messages
from django.contrib.auth.models import User
from django.core.files.storage import default_storage
from django.core.mail import EmailMultiAlternatives
from django.http import StreamingHttpResponse, HttpResponseRedirect
from django.urls import reverse
from django.db.models import Max, Q
from django.shortcuts import get_object_or_404, redirect, render
from django.template import Context
from django.template.loader import get_template
from django.utils import timezone
from django.utils.encoding import smart_text
from docx import Document

from pymarc import Record, Field, record_to_xml

from core import email, models, log, forms
from core.cache import cache_result
from core.decorators import is_reviewer
from core.files import handle_marc21_file
from core.util import (
    get_setting,
    strip_html_tags,
)

from editorialreview import models as editorialreview_models
from manager import forms as manager_forms
from review import forms as review_forms, models as review_models, \
    logic as review_logic
from revisions import models as revisions_models
from submission import logic as submission_logic, models as submission_models


def record_field(tag, indicators, subfields):
    return Field(tag=tag, indicators=indicators, subfields=subfields)


def record_control_field(tag, field):
    return Field(tag=tag, data=field)


def book_to_mark21_file(book, owner, xml=False):
    """
    Number and value explanation:
    http://www.loc.gov/marc/bibliographic/bdleader.html
    Adding Leader tags.
    """
    record = Record()  # New record.

    leader_list = list(record.leader)
    leader_list[5] = 'n'  # New
    leader_list[6] = 'a'  # For manuscript file use 't'
    leader_list[7] = 'm'  # Monograph
    leader_list[9] = 'a'
    leader_list[19] = '#'
    record.leader = "".join(leader_list)

    # Category of material  - Text
    record.add_field(record_control_field('007', 't'))

    # Languages
    languages = book.languages.all()
    if languages:
        for lang in languages:
            record.add_field(record_control_field('008', lang.code))
    else:
        record.add_field(record_control_field('008', 'eng'))

    # ISBN - International Standard Book Number.
    isbn = models.Identifier.objects.filter(
        book=book
    ).exclude(
        identifier='pub_id'
    ).exclude(
        identifier='urn'
    ).exclude(
        identifier='doi'
    )

    for identifier in isbn:
        if book.book_type:
            record.add_field(
                record_field(
                    '020',
                    ['#', '#'],
                    ['a', str(identifier.value) + ' ' + book.book_type]
                )
            )
        else:
            record.add_field(
                record_field(
                    '020',
                    ['#', '#'],
                    ['a', str(identifier.value)]
                )
            )

    # Source of acquisition.
    base_url = get_setting('base_url', 'general')

    book_url = 'http://%s/editor/submission/%s/' % (base_url, book.id)
    record.add_field(record_field('030', ['#', '#'], ['b', book_url]))

    authors = book.author.all()  # Main entry - Personal name.
    author_names = ''

    for author in authors:
        author_names = author_names + author.full_name() + ' '
        name = author.last_name + ', ' + author.first_name

        if author.middle_name:
            name = name + ' ' + author.middle_name[:1] + '.'

        record.add_field(record_field('100', ['1', '#'], ['a', name]))

    title_words = book.title.split(' ')  # Title statement.
    first_word = title_words[0]

    if first_word.lower() == 'the':
        record.add_field(
            record_field(
                '245',
                ['1', '4'],
                ['a', book.title, 'c', author_names]
            )
        )
    else:
        record.add_field(
            record_field(
                '245',
                ['1', '0'],
                ['a', book.title, 'c', author_names]
            )
        )

    # Publication.
    press_name = get_setting('press_name', 'general')
    city = get_setting('city', 'general')

    publication_info = []
    if book.publication_date:  # Press' city.
        if city:
            publication_info.append('a')
            publication_info.append(city)
        if press_name:  # Press' name.
            publication_info.append('b')
            publication_info.append(press_name)
        publication_info.append('c')  # Date of Publication.
        publication_info.append(str(book.publication_date))
        record.add_field(record_field('260', ['#', '#'], publication_info))

    if book.pages:  # Physical details.
        record.add_field(
            record_field('300', ['#', '#'], ['a', str(book.pages) + ' pages'])
        )

    record.add_field(  # Content type.
        record_field('336', ['#', '#'], ['a', 'text', '2', 'rdacontent'])
    )

    record.add_field(  # Media type.
        record_field('337', ['#', '#'], ['a', 'unmediated', '2', 'rdamedia'])
    )

    record.add_field(  # Carrier type.
        record_field('338', ['#', '#'], ['a', 'volume', '2', 'rdacarrier'])
    )

    if languages:  # Language note.
        for lang in languages:
            record.add_field(
                record_field('546', ['#', '#'], ['a', lang.display]))
    else:
        record.add_field(record_field('546', ['#', '#'], ['a', 'In English']))

    press_editors = book.press_editors.all()

    for editor in press_editors:  # Editors.
        record.add_field(record_field(
            '700',
            ['1', '#'],
            ['a', '%s, %s' % (editor.last_name,
                              editor.first_name), 'e', 'Press editor']))

    if book.series:  # Series.
        record.add_field(
            record_field('830', ['#', '0'], ['a', book.series.title]))
        if book.series.editor:
            record.add_field(record_field(
                '700',
                ['1', '#'],
                ['a', '%s, %s' % (book.series.editor.last_name,
                                  book.series.editor.first_name),
                 'e', 'Series editor']
            ))

    title = book.title  # Add record to file.
    if not xml:
        filename = 'book_' + str(book.id) + '_' + re.sub(
            r'[^a-zA-Z0-9\n\.]', '', title.lower()
        ) + '_marc21.dat'
        _file = handle_marc21_file(record.as_marc(), filename, book, owner)
    else:
        filename = 'book_' + str(book.id) + '_' + re.sub(
            r'[^a-zA-Z0-9\n\.]', '', title.lower()
        ) + '_marc21.xml'
        content = record_to_xml(record, quiet=False, namespace=False)
        _file = handle_marc21_file(content, filename, book, owner)

    return _file.pk
    # add handle_file ?


def book_to_mark21_file_download_content(book, owner, content, xml=False):
    title = book.title

    if not content or content.isspace():
        content = 'No content found.'

    if not xml:
        filename = 'book_' + str(book.id) + '_' + re.sub(
            r'[^a-zA-Z0-9\n\.]', '', title.lower()
        ) + '_marc21.dat'
        _file = handle_marc21_file(content, filename, book, owner)
    else:
        filename = 'book_' + str(book.id) + '_' + re.sub(
            r'[^a-zA-Z0-9\n\.]', '', title.lower()
        ) + '_marc21.xml'
        _file = handle_marc21_file(content, filename, book, owner)

    return _file.pk


def book_to_mark21_file_content(book, owner, xml=False):
    """ Number and value explanation:
    http://www.loc.gov/marc/bibliographic/bdleader.html
    """
    record = Record()  # New record.

    leader_list = list(record.leader)  # Adding Leader tags.
    leader_list[5] = 'n'  # New
    leader_list[6] = 'a'  # For manuscript file use 't'
    leader_list[7] = 'm'  # Monograph
    leader_list[9] = 'a'
    leader_list[19] = '#'
    record.leader = "".join(leader_list)

    # Category of material  - Text
    record.add_field(record_control_field('007', 't'))

    languages = book.languages.all()  # Languages.
    if languages:
        for lang in languages:
            record.add_field(record_control_field('008', lang.code))
    else:
        record.add_field(record_control_field('008', 'eng'))

    isbn = models.Identifier.objects.filter(
        # International Standard Book Number
        book=book
    ).exclude(
        identifier='pub_id'
    ).exclude(
        identifier='urn'
    ).exclude(
        identifier='doi'
    )

    for identifier in isbn:
        if book.book_type:
            record.add_field(record_field(
                '020',
                ['#', '#'],
                ['a', str(identifier.value) + ' ' + book.book_type]
            ))
        else:
            record.add_field(
                record_field('020', ['#', '#'], ['a', str(identifier.value)]))

    # Source of acquisition.
    base_url = get_setting('base_url', 'general', default='localhost:8000')
    book_url = 'http://%s/editor/submission/%s/' % (base_url, book.id)
    record.add_field(record_field('030', ['#', '#'], ['b', book_url]))

    authors = book.author.all()  # Main entry - Personal name.
    author_names = ''

    for author in authors:
        author_names = author_names + author.full_name() + ' '
        name = author.last_name + ', ' + author.first_name

        if author.middle_name:
            name = name + ' ' + author.middle_name[:1] + '.'

        record.add_field(record_field('100', ['1', '#'], ['a', name]))

    title_words = book.title.split(' ')  # Title statement.
    first_word = title_words[0]

    if first_word.lower() == 'the':
        record.add_field(record_field(
            '245',
            ['1', '4'],
            ['a', book.title, 'c', author_names]
        ))
    else:
        record.add_field(record_field(
            '245',
            ['1', '0'],
            ['a', book.title, 'c', author_names]
        ))

    # Publication.
    press_name = get_setting('press_name', 'general')
    city = get_setting('city', 'general')
    publication_info = []

    if book.publication_date:
        if city:  # Press' city.
            publication_info.append('a')
            publication_info.append(str(city))
        if press_name:  # Press' name.
            publication_info.append('b')
            publication_info.append(str(press_name))

        publication_info.append('c')  # Date of Publication.
        publication_info.append(str(book.publication_date))
        record.add_field(record_field('260', ['#', '#'], publication_info))

    if book.pages:  # Physical details.
        record.add_field(record_field(
            '300',
            ['#', '#'],
            ['a', str(book.pages) + ' pages']
        ))

    record.add_field(record_field(  # Content type.
        '336',
        ['#', '#'],
        ['a', 'text', '2', 'rdacontent']
    ))

    record.add_field(record_field(  # Media type.
        '337',
        ['#', '#'],
        ['a', 'unmediated', '2', 'rdamedia']
    ))

    record.add_field(record_field(  # Carrier type.
        '338',
        ['#', '#'],
        ['a', 'volume', '2', 'rdacarrier']
    ))

    if languages:  # Language note.
        for lang in languages:
            record.add_field(
                record_field('546', ['#', '#'], ['a', lang.display])
            )
    else:
        record.add_field(record_field('546', ['#', '#'], ['a', 'In English']))

    press_editors = book.press_editors.all()

    for editor in press_editors:  # Editors.
        record.add_field(record_field(
            '700',
            ['1', '#'],
            [
                'a',
                '%s, %s' % (
                    editor.last_name,
                    editor.first_name
                ),
                'e',
                'Press editor'
            ]
        ))

    if book.series:  # Series.
        record.add_field(
            record_field('830', ['#', '0'], ['a', book.series.title]))
        if book.series.editor:
            record.add_field(record_field(
                '700',
                ['1', '#'],
                [
                    'a', '%s, %s' % (
                        book.series.editor.last_name,
                        book.series.editor.first_name
                    ),
                    'e',
                    'Series editor'
                ]
            ))

    title = book.title  # Add record to file.

    if not xml:
        filename = 'book_' + str(book.id) + '_' + re.sub(
            r'[^a-zA-Z0-9\n\.]', '', title.lower()
        ) + '_marc21.dat'
        handle_marc21_file(record.as_marc(), filename, book, owner)
        content = record.as_marc()
    else:
        filename = 'book_' + str(book.id) + '_' + re.sub(
            r'[^a-zA-Z0-9\n\.]', '', title.lower()
        ) + '_marc21.xml'
        content = record_to_xml(record, quiet=False, namespace=False)
        handle_marc21_file(content, filename, book, owner)

    return content or None


def get_author_emails(submission_id, term):
    submission = get_object_or_404(models.Book, pk=submission_id)
    authors = submission.author.all()
    results = []

    for author in authors:
        name = author.full_name()
        author_json = {
            'id': author.id,
            'label': author.full_name(),
            'value': author.author_email
        }

        if term:
            if term.lower() in name.lower():
                results.append(author_json)

    return results


def get_editor_emails(submission_id, term):
    submission = get_object_or_404(models.Book, pk=submission_id)
    editors = get_editors(submission)
    results = []

    for editor in editors:
        if hasattr(editor, 'full_name'):
            name = editor.full_name()
        else:
            name = '{first_name} {last_name}'.format(
                first_name=editor.first_name,
                last_name=editor.last_name
            )

        editor_json = {'id': editor.id, 'label': name}

        if hasattr(editor, 'author_email'):
            editor_json['value'] = editor.author_email
        else:
            editor_json['value'] = editor.email

        if term:
            if term.lower() in name.lower():
                results.append(editor_json)

    return results


def get_all_user_emails(term):
    users = User.objects.all()
    results = []

    for user in users:
        if user.profile and hasattr(user.profile, 'full_name'):
            name = user.profile.full_name()
        else:
            name = '{first_name} {last_name}'.format(
                first_name=user.first_name,
                last_name=user.last_name
            )

        user_json = {'id': user.id, 'label': name, 'value': user.email}

        if term:
            if term.lower() in name.lower():
                results.append(user_json)

    return results


def get_onetasker_emails(submission_id, term):
    submission = get_object_or_404(models.Book, pk=submission_id)
    onetaskers = submission.get_onetaskers()

    results = []
    for user in onetaskers:
        user_json = {}
        name = '{first_name} {last_name}'.format(
                first_name=user.first_name,
                last_name=user.last_name
            )
        user_json['id'] = user.id
        user_json['label'] = user.profile.full_name()
        user_json['value'] = user.email
        if (
            not string_any(user_json['value'] for _ in results) and
            term.lower() in name.lower()
        ):
            results.append(user_json)

    return results


def get_proposal_emails(proposal_id, term):
    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    results = []
    user = proposal.owner
    name = '{first_name} {last_name}'.format(
        first_name=user.first_name,
        last_name=user.last_name
    )
    user_json = {
        'id': user.id,
        'label': user.profile.full_name(),
        'value': user.email,
    }

    if (
            not string_any(user_json['value'] for _ in results) and
            term.lower() in name.lower()
    ):
        results.append(user_json)

    if proposal.requestor:
        user = proposal.requestor
        name = '{first_name} {last_name}'.format(
                first_name=user.first_name,
                last_name=user.last_name
            )
        user_json = {
            'id': user.id,
            'label': user.profile.full_name(),
            'value': user.email,
        }
        if (
                not string_any(user_json['value'] for result in results) and
                term.lower() in name.lower()
        ):
            results.append(user_json)

    return results


def get_editors(book):
    press_editors = book.press_editors.all()
    book_editors = book.book_editors.all()

    if book.series:
        series_editor = book.series.editor

        if series_editor:
            series_editor_list = [series_editor]
            press_editor_list = [
                editor for editor in press_editors
                if not editor == series_editor_list[0]
            ]
        else:
            series_editor_list = []
            press_editor_list = [editor for editor in press_editors]
    else:
        series_editor_list = []
        press_editor_list = [editor for editor in press_editors]

    if book_editors:
        book_editor_list = [
            editor for editor in book_editors
            if editor not in press_editor_list
        ]
    else:
        book_editor_list = []

    return press_editor_list + series_editor_list + book_editor_list


def clean_email_list(addresses):
    list_of_email_addresses = []

    for address in addresses:
        if '@' in address:
            if address.replace(" ", "") not in list_of_email_addresses:
                list_of_email_addresses.append(address.replace(" ", ""))

    if len(list_of_email_addresses) < 1:
        return None

    return list_of_email_addresses


def send_email(
        subject,
        context,
        from_email,
        to,
        html_template,
        text_template=None,
):
    plaintext = get_template(text_template)
    htmly = get_template(html_template)

    con = Context(context)
    text_content = plaintext.render(con)
    html_content = htmly.render(con)
    msg = EmailMultiAlternatives(subject, text_content, from_email, [to])
    msg.attach_alternative(html_content, "text/html")
    msg.send()


@cache_result(300)
def press_settings():
    _dict = {}

    for group in models.SettingGroup.objects.all():
        _dict[group.name] = {
            setting.name: setting.value for setting in
            models.Setting.objects.filter(group=group)
        }

    return _dict


def task_count(request):
    """
    Counts the number of incomplete tasks assigned to the user of a request
    :param request: the request containing the user object used in the query
    :return: the number of incomplete talks assigned to the request's user
    """
    if request.user.is_authenticated:
        return models.Task.objects.filter(
            assignee=request.user,
            completed__isnull=True,
        ).count()
    else:
        return 0


def review_assignment_count(request):
    """
    Counts the number of active reviews assigned to the user of a request
    :param request: the request containing the user object used in the query
    :return: the number of active reviews assigned to the request's user
    """
    if request.user.is_authenticated:
        return models.ReviewAssignment.objects.filter(
            user=request.user,
            completed__isnull=True,
            declined__isnull=True,
            withdrawn=False
        ).count() + submission_models.ProposalReview.objects.filter(
            user=request.user,
            completed__isnull=True,
            declined__isnull=True,
            withdrawn=False
        ).count() + models.ReviewAssignment.objects.filter(
            user=request.user,
            completed__isnull=False,
            declined__isnull=True,
            reopened=True
        ).count() + editorialreview_models.EditorialReview.objects.filter(
            user=request.user,
            completed__isnull=True,
            withdrawn=False
        ).count()
    else:
        return 0


def author_tasks(user):
    """
    Returns a list of revision, typeset and copyedit tasks for a given user
    :param user: the user for whom to get tasks
    :return: a list of incomplete assigned to the user
    """
    task_list = []

    if user.is_authenticated:
        base_url = get_setting('base_url', 'general')
        revision_tasks = revisions_models.Revision.objects.filter(
            book__owner=user,
            requested__isnull=False,
            completed__isnull=True
        ).select_related(
            'book'
        )
        copyedit_tasks = models.CopyeditAssignment.objects.filter(
            book__owner=user,
            author_invited__isnull=False,
            author_completed__isnull=True
        ).select_related(
            'book'
        )
        typeset_tasks = models.TypesetAssignment.objects.filter(
            book__owner=user,
            author_invited__isnull=False,
            author_completed__isnull=True
        ).select_related(
            'book'
        )

        for revision in revision_tasks:
            task_list.append(
                {
                    'task': 'Revisions Requested',
                    'title': revision.book.title,
                    'url': 'http://%s/revisions/%s' % (base_url, revision.id)
                }
            )

        for copyedit in copyedit_tasks:
            task_list.append(
                {
                    'task': 'Copyedit Review',
                    'title': copyedit.book.title,
                    'url': 'http://%s/copyedit/book/%s/edit/%s/author/' % (
                        base_url,
                        copyedit.book.id,
                        copyedit.id
                    )
                }
            )

        for typeset in typeset_tasks:
            task_list.append(
                {
                    'task': 'Typesetting Review',
                    'title': typeset.book.title,
                    'url': 'http://%s/typeset/book/%s/typeset/%s/author/' % (
                        base_url,
                        typeset.book.id,
                        typeset.id
                    )
                }
            )

    return task_list


def typesetter_tasks(user):
    active = models.TypesetAssignment.objects.filter(
        (
                Q(requested__isnull=False) &
                Q(completed__isnull=True)
        ) | (
                Q(typesetter_invited__isnull=False) &
                Q(typesetter_completed__isnull=True)
        ),
        typesetter=user,
    ).select_related(
        'book',
    ).exclude(
        declined__isnull=False,
    )

    completed = models.TypesetAssignment.objects.filter(
        (
                Q(completed__isnull=False) &
                Q(typesetter_completed__isnull=True)
        ) | (
                Q(completed__isnull=False) &
                Q(typesetter_completed__isnull=False)
        ),
        typesetter=user,
    ).select_related(
        'book',
    ).order_by(
        '-completed',
    )[:5]

    return {'active': active, 'completed': completed}


def copyeditor_tasks(user):
    active = models.CopyeditAssignment.objects.filter(
        copyeditor=user,
        completed__isnull=True,
    ).exclude(
        declined__isnull=False,
    ).select_related(
        'book',
    )

    completed = models.CopyeditAssignment.objects.filter(
        copyeditor=user,
        completed__isnull=False
    ).select_related(
        'book'
    ).order_by(
        '-completed'
    )[:5]

    return {'active': active, 'completed': completed}


def indexer_tasks(user):
    active = models.IndexAssignment.objects.filter(
        indexer=user,
        completed__isnull=True
    ).exclude(
        declined__isnull=False
    ).select_related(
        'book'
    )

    completed = models.IndexAssignment.objects.filter(
        indexer=user,
        completed__isnull=False
    ).select_related(
        'book'
    ).order_by(
        '-completed'
    )[:5]

    return {'active': active, 'completed': completed}


def onetasker_tasks(user):
    active = []
    completed = []

    active_copyeditor_tasks = copyeditor_tasks(user).get('active')
    completed_copyeditor_tasks = copyeditor_tasks(user).get('completed')
    active_typesetter_tasks = typesetter_tasks(user).get('active')
    completed_typesetter_tasks = typesetter_tasks(user).get('completed')
    active_indexer_tasks = indexer_tasks(user).get('active')
    completed_indexer_tasks = indexer_tasks(user).get('completed')

    for assignment in active_copyeditor_tasks:
        active.append({'assignment': assignment, 'type': 'copyedit', })

    for assignment in active_typesetter_tasks:
        active.append({'assignment': assignment, 'type': 'typesetting'})

    for assignment in active_indexer_tasks:
        active.append({'assignment': assignment, 'type': 'indexing'})

    for assignment in completed_copyeditor_tasks:
        completed.append({'assignment': assignment, 'type': 'copyedit'})

    for assignment in completed_typesetter_tasks:
        completed.append({'assignment': assignment, 'type': 'typesetting'})

    for assignment in completed_indexer_tasks:
        completed.append({'assignment': assignment, 'type': 'indexing'})

    return {'completed': completed, 'active': active}


def create_new_review_round(book):
    latest_round = models.ReviewRound.objects.filter(
        book=book
    ).aggregate(
        max=Max('round_number')
    )
    next_round = (
        latest_round.get('max') + 1 if
        latest_round.get('max') and
        latest_round.get('max') > 0
        else 1
    )

    return models.ReviewRound.objects.create(book=book, round_number=next_round)


def build_time_line_editing_copyedit(copyedit):
    timeline = []
    overdue = False

    if copyedit.accepted:
        if copyedit.completed and copyedit.completed > copyedit.due:
            overdue = True

        timeline.append({
            'stage': 'Requested',
            'date': copyedit.requested,
            'overdue': overdue,
        })
        timeline.append({
            'stage': 'Accepted',
            'date': copyedit.accepted,
            'overdue': overdue,
        })

        if copyedit.completed:
            if overdue:
                timeline.append({
                    'stage': 'Completed',
                    'date': copyedit.completed,
                    'overdue': overdue,
                })
            else:
                timeline.append({
                    'stage': 'Completed',
                    'date': copyedit.completed,
                    'overdue': overdue,
                })
        else:
            timeline.append({
                'stage': 'Due',
                'date': copyedit.due,
                'overdue': overdue,
            })
        timeline.append({
            'stage': 'Editor Review',
            'date': copyedit.editor_review,
            'overdue': overdue,
        })
        timeline.append({
            'stage': 'Author Invited',
            'date': copyedit.author_invited,
            'overdue': overdue,
        })
        timeline.append({
            'stage': 'Author completed',
            'date': copyedit.author_completed,
            'overdue': overdue,
        })
    else:
        timeline.append({
            'stage': 'Requested',
            'date': copyedit.requested,
            'overdue': overdue,
        })
        timeline.append({
            'stage': 'Declined',
            'date': copyedit.declined,
            'declined': True,
        })

    clean_timeline = []

    for time in timeline:
        if time['date']:
            if isinstance(time['date'], datetime.datetime):
                time['date'] = time['date'].date()
            clean_timeline.append(time)

    return sorted(clean_timeline, key=lambda k: k['date'])


def build_time_line_editing_indexer(index):
    timeline = []
    overdue = False

    if index.accepted:
        if index.completed and index.completed > index.due:
            overdue = True
        timeline.append({
            'stage': 'Requested',
            'date': index.requested,
            'overdue': overdue,
        })
        timeline.append({
            'stage': 'Accepted',
            'date': index.accepted,
            'overdue': overdue,
        })
        if index.completed:
            if overdue:
                timeline.append({
                    'stage': 'Due',
                    'date': index.due,
                    'overdue': overdue,
                })
                timeline.append({
                    'stage': 'Completed',
                    'date': index.completed,
                    'overdue': overdue,
                })
            else:
                timeline.append({
                    'stage': 'Completed',
                    'date': index.completed,
                    'overdue': overdue,
                })
                timeline.append({
                    'stage': 'Due',
                    'date': index.due,
                    'overdue': overdue,
                })
        else:
            timeline.append({
                'stage': 'Due',
                'date': index.due,
                'overdue': overdue,
            })
    else:
        timeline.append({
            'stage': 'Declined',
            'date': index.declined,
            'declined': True,
        })
        timeline.append({'stage': 'Due', 'date': index.due, 'overdue': overdue})

    clean_timeline = []

    for time in timeline:
        if time['date']:
            if isinstance(time['date'], datetime.datetime):
                time['date'] = time['date'].date()
            clean_timeline.append(time)

    return sorted(clean_timeline, key=lambda k: k['date'])


def build_time_line(book):
    timeline = []

    if book.stage:
        timeline.append({
            'stage': 'Declined',
            'date': book.stage.declined,
        })
        timeline.append({
            'stage': 'Publication',
            'date': book.stage.publication,
        })
        timeline.append({
            'stage': 'Production',
            'date': book.stage.production,
        })
        timeline.append({
            'stage': 'Typesetting',
            'date': book.stage.typesetting,
        })
        timeline.append({
            'stage': 'Indexing',
            'date': book.stage.indexing,
        })
        timeline.append({
            'stage': 'Copyediting',
            'date': book.stage.copyediting,
        })
        timeline.append({
            'stage': 'Editing',
            'date': book.stage.editing,
        })
        timeline.append({
            'stage': 'External Review',
            'date': book.stage.external_review,
        })
        timeline.append({
            'stage': 'Internal Review',
            'date': book.stage.internal_review,
        })
        timeline.append({
            'stage': 'Review',
            'date': book.stage.review,
        })

        if book.proposal:
            timeline.append({
                'stage': 'Proposal Submitted',
                'date': book.proposal.date_submitted,
            })
            timeline.append({
                'stage': 'Proposal Review Started',
                'date': book.proposal.date_review_started,
            })
            timeline.append({
                'stage': 'Proposal Accepted',
                'date': book.proposal.date_accepted,
            })

        timeline.append({
            'stage': 'Book Submitted',
            'date': book.stage.submission,
        })
        timeline.append({
            'stage': 'Proposal',
            'date': book.stage.proposal,
        })

    clean_timeline = []
    for time in timeline:
        if time['date']:
            if isinstance(time['date'], datetime.datetime):
                time['date'] = time['date'].date()
            clean_timeline.append(time)

    return sorted(clean_timeline, key=lambda k: k['date'], reverse=True)


def send_proposal_review_request(
        request,
        proposal,
        review_assignment,
        email_text,
        attachment=None,
        access_key=None,
):
    from_email = request.user.email or get_setting('from_address', 'email')
    base_url = get_setting('base_url', 'general')
    press_name = get_setting('press_name', 'general')

    if access_key:
        review_url = "http://{0}{1}".format(
            base_url,
            reverse(
                'view_proposal_review_decision_access_key',
                kwargs={
                    'proposal_id': proposal.id,
                    'assignment_id': review_assignment.id,
                    'access_key': access_key,
                }
            )
        )
    else:
        review_url = "http://{0}{1}".format(
            base_url,
            reverse(
                'view_proposal_review_decision',
                kwargs={
                    'proposal_id': proposal.id,
                    'assignment_id': review_assignment.id,
                }
            )
        )

    if request:
        from_email = "%s <%s>" % (
            request.user.profile.full_name(),
            from_email,
        )

    context = {
        'review': review_assignment,
        'review_url': review_url,
        'proposal': proposal,
        'press_name': press_name,
    }
    email.send_email(
        get_setting(
            'proposal_review_request_subject',
            'email_subject',
            'Proposal Review Request'
        ),
        context,
        from_email,
        review_assignment.user.email,
        email_text,
        proposal=proposal,
        attachment=attachment,
        request=request,
        kind='proposal_review',
        access_key=access_key,
    )


def send_proposal_review_reopen_request(
        request,
        proposal,
        review_assignment,
        email_text,
        attachment=None,
):
    from_email = request.user.email or get_setting('from_address', 'email')
    base_url = get_setting('base_url', 'general')
    press_name = get_setting('press_name', 'general')

    if request:
        from_email = "%s <%s>" % (
            request.user.profile.full_name(),
            from_email,
        )

    review_url = "http://{0}{1}".format(
        base_url,
        reverse(
            'view_proposal_review_decision',
            kwargs={
                'proposal_id': proposal.id,
                'assignment_id': review_assignment.id
            }
        ))

    context = {
        'review': review_assignment,
        'review_url': review_url,
        'proposal': proposal,
        'press_name': press_name,
    }
    email.send_email(
        get_setting(
            'proposal_review_reopen_subject',
            'email_subject',
            'Proposal Review Assignment has reopened'
        ),
        context,
        from_email,
        review_assignment.user.email,
        email_text,
        proposal=proposal,
        attachment=attachment,
        request=request,
        kind='proposal_review',
    )


def order_data(data, relations):
    ordered_data = []
    for relation in relations:
        if relation.element.name in data:
            ordered_data.append(
                [relation.element.name, data[relation.element.name]]
            )

    return ordered_data


def decode_json(json_data):
    return json.loads(json_data)


def encode_data(data):
    return smart_text(json.dumps(data))


def close_active_reviews(proposal):
    for review in proposal.review_assignments.all():
        review.completed = timezone.now()
        review.save()


def create_submission_from_proposal(proposal, proposal_type):
    book = models.Book(
        title=proposal.title,
        subtitle=proposal.subtitle,
        owner=proposal.owner,
        book_type=proposal_type,
        submission_stage=1,
    )

    book.save()

    if book.book_type == 'monograph':
        submission_logic.copy_author_to_submission(proposal.owner, book)
    elif book.book_type == 'edited_volume':
        submission_logic.copy_editor_to_submission(proposal.owner, book)

    book.save()

    return book


def handle_typeset_assignment(
        book,
        typesetter,
        files,
        due_date,
        email_text,
        requestor,
        attachment,
):
    new_typesetter = models.TypesetAssignment(
        book=book,
        typesetter=typesetter,
        requestor=requestor,
        due=due_date,
    )

    new_typesetter.save()

    for _file in files:
        new_typesetter.files.add(_file)

    new_typesetter.save()
    send_invite_typesetter(
        book,
        new_typesetter,
        email_text,
        requestor,
        attachment,
    )

    log.add_log_entry(
        book=book,
        user=requestor,
        kind='typeser',
        message='Typesetter %s %s assigned. Due %s' % (
            typesetter.first_name,
            typesetter.last_name, due_date
        ),
        short_name='Typeset Assignment',
    )


def send_decision_ack(
        request,
        book,
        decision,
        email_text,
        url=None,
        attachment=None
):
    """ Email Handlers - TODO: move to email.py? """
    from_email = request.user.email or get_setting('from_address', 'email')

    if not decision == 'decline':
        decision_full = "Move to " + decision
    else:
        decision_full = 'Reject Submission'

    authors = book.author.all()

    for author in authors:
        context = {
            'submission': book,
            'author': author,
            'decision': decision,
            'link_to_page': url,
        }
        kind = "submission"
        subject = get_setting(
            'submission_decision_update_subject',
            'email_subject',
            'Submission decision update: %s' % decision_full
        )

        if attachment:

            email.send_email(
                subject,
                context,
                from_email,
                author.author_email,
                email_text,
                kind=kind,
                book=book,
                attachment=attachment,
            )
        else:
            email.send_email(
                subject,
                context,
                from_email,
                author.author_email,
                email_text,
                kind=kind,
                book=book,
            )


def send_editorial_decision_ack(
        request,
        review_assignment,
        contact,
        decision,
        email_text,
        url=None,
        attachment=None,
):
    from_email = request.user.email or get_setting('from_address', 'email')
    publishing_committee = get_setting('publishing_committee', 'general')

    decision_full = decision

    if contact == 'editorial-board':
        editors = review_assignment.editorial_board.all()

        for editor in editors:
            context = {
                'submission': review_assignment.book,
                'editor': editor.profile.full_name(),
                'decision': decision,
                'link_to_page': url,
            }
            subject = get_setting(
                'submission_decision_update_subject',
                'email_subject',
                'Submission decision update: %s' % decision_full,
            )
            email.send_email(
                subject,
                context,
                from_email,
                editor.email,
                email_text,
                book=review_assignment.book,
                attachment=attachment,
                kind='submission',
            )
    elif contact == 'author':
        authors = review_assignment.book.author.all()

        for author in authors:
            context = {
                'submission': review_assignment.book,
                'name': author.full_name(),
                'decision': decision,
                'link_to_page': url,
            }
            subject = get_setting(
                'submission_decision_update_subject',
                'email_subject',
                'Submission decision update: %s' % decision_full,
            )
            email.send_email(
                subject,
                context,
                from_email,
                author.author_email,
                email_text,
                book=review_assignment.book,
                attachment=attachment,
                kind='submission',
            )
    elif contact == 'publishing-committee':
        emails = clean_email_list(publishing_committee.split(';'))
        context = {
            'submission': review_assignment.book,
            'name': 'Publishing Committee',
            'decision': decision,
            'link_to_page': url,
        }

        for current_email in emails:
            subject = get_setting(
                'submission_decision_update_subject',
                'email_subject',
                'Submission decision update: %s' % decision_full,
            )
            email.send_email(
                subject,
                context,
                from_email,
                current_email,
                email_text,
                book=review_assignment.book,
                attachment=attachment,
                kind='submission',
            )


def send_production_editor_ack(
        request,
        book,
        editor,
        email_text,
        attachment=None
):
    """ Email Handlers - TODO: move to email.py? """
    from_email = request.user.email or get_setting('from_address', 'email')
    context = {'submission': book, 'editor': editor}
    subject = get_setting(
        'production_editor_subject',
        'email_subject',
        'Production Editor for {0}'.format(book.full_title),
    )

    email.send_email(
        subject,
        context,
        from_email,
        editor.email,
        email_text,
        book=book,
        attachment=attachment,
        kind='production',
    )


def send_review_request(
        request,
        book,
        review_assignment,
        email_text,
        sender,
        attachment=None,
        access_key=None,
):
    from_email = request.user.email or get_setting('from_address', 'email')
    base_url = get_setting('base_url', 'general')
    press_name = get_setting('press_name', 'general')

    if access_key:
        decision_url = (
            'http://{base_url}/review/{review_type}/{book_id}/assignment/'
            '{review_assignment_id}/access_key/{access_key}/decision/'.format(
                base_url=base_url,
                review_type=review_assignment.review_type,
                book_id=book.id,
                review_assignment_id=review_assignment.id,
                access_key=access_key,
            )
        )
    else:
        decision_url = (
            'http://{base_url}/review/{review_type}/{book_id}/'
            'assignment/{review_assignment_id}/decision/'.format(
                base_url=base_url,
                review_type=review_assignment.review_type,
                book_id=book.id,
                review_assignment_id=review_assignment.id,
            )
        )

    context = {
        'book': book,
        'review': review_assignment,
        'decision_url': decision_url,
        'sender': sender,
        'base_url': base_url,
        'press_name': press_name,
    }

    email.send_email(
        subject=get_setting(
            'review_request_subject',
            'email_subject',
            'Review Request',
        ),
        context=context,
        from_email=from_email,
        to=review_assignment.user.email,
        html_template=email_text,
        book=book,
        attachment=attachment,
        kind='review',
    )


def send_proposal_book_editor(
        request,
        proposal,
        email_text,
        sender,
        editor_email,
):
    from_email = request.user.email or get_setting('from_address', 'email')
    if request:
        from_email = "%s <%s>" % (
            request.user.profile.full_name(),
            from_email,
        )

    context = {'proposal': proposal, 'sender': sender}
    subject = get_setting(
        'proposal_book_editors_subject',
        'email_subject',
        '[abp] Proposal Book Editors: Update',
    )
    email.send_email(
        subject,
        context,
        from_email,
        editor_email,
        email_text,
        proposal=proposal,
        request=request,
        kind='proposal',
    )


def send_proposal_decline(request, proposal, email_text, sender):
    from_email = request.user.email or get_setting('from_address', 'email')
    if request:
        from_email = "%s <%s>" % (
            request.user.profile.full_name(),
            from_email,
        )

    context = {'proposal': proposal, 'sender': sender}
    subject = get_setting(
        'proposal_declined_subject',
        'email_subject',
        '[abp] Proposal Declined',
    )
    email.send_email(
        subject,
        context,
        from_email,
        proposal.owner.email,
        email_text,
        proposal=proposal,
        request=request,
        kind='proposal',
    )


def send_proposal_update(
        request,
        proposal,
        email_text,
        sender,
        receiver
):
    from_email = request.user.email or get_setting('from_address', 'email')
    context = {'proposal': proposal, 'sender': sender, 'receiver': receiver}
    subject = get_setting(
        'proposal_update_subject',
        'email_subject',
        '[abp] Proposal Update',
    )
    email.send_email(
        subject,
        context,
        from_email,
        proposal.owner.email,
        email_text,
        proposal=proposal,
        kind='proposal',
    )


def send_proposal_submission_ack(request, proposal, email_text, owner):
    from_email = request.user.email or get_setting('from_address', 'email')
    press_name = get_setting('press_name', 'general')
    principal_contact_name = get_setting('primary_contact_name', 'general')

    context = {
        'proposal': proposal,
        'owner': owner,
        'press_name': press_name,
        'principal_contact_name': principal_contact_name,
    }
    subject = get_setting(
        'proposal_submission_ack_subject',
        'email_subject',
        '[abp] Proposal Submission Acknowledgement',
    )
    email.send_email(
        subject,
        context,
        from_email,
        proposal.owner.email,
        email_text,
        proposal=proposal,
        kind='proposal',
    )


def send_proposal_change_owner_ack(request, proposal, email_text, owner):
    from_email = request.user.email or get_setting('from_address', 'email')
    press_name = get_setting('press_name', 'general')
    principal_contact_name = get_setting('primary_contact_name', 'general')

    context = {
        'proposal': proposal,
        'receiver': owner,
        'sender': request.user,
        'base_url': get_setting('base_url', 'general'),
        'press_name': press_name,
        'principal_contact_name': principal_contact_name,
    }
    subject = get_setting(
        'change_principal_contact_proposal_subject',
        'email_subject',
        '[abp] Proposal Owner Change',
    )
    email.send_email(
        subject,
        context,
        from_email,
        proposal.owner.email,
        email_text,
        proposal=proposal,
        kind='proposal',
        request=request,
    )


def send_task_decline(assignment, _type, email_text, sender, request):
    if request.user.is_authenticated:
        from_email = request.user.email
        from_email = "%s <%s>" % (
            request.user.profile.full_name(),
            from_email,
        )
    else:
        from_email = get_setting('from_address', 'email')

    context = {'assignment': assignment, 'sender': sender}
    subject = get_setting(
        'assignment_declined_subject',
        'email_subject',
        '[abp] %s Assignment [id<%s>] Declined') % (
                  _type.title(),
                  assignment.id,
              )
    email.send_email(
        subject,
        context,
        from_email,
        assignment.requestor.email,
        email_text,
        request=request,
        kind='workflow',
    )


def send_proposal_accept(
        request,
        proposal,
        email_text,
        submission, sender,
        attachment=None,
):
    from_email = request.user.email or get_setting('from_address', 'email')

    if request:
        from_email = "%s <%s>" % (
            request.user.profile.full_name(),
            from_email,
        )

    context = {
        'base_url': get_setting('base_url', 'general'),
        'proposal': proposal,
        'submission': submission,
        'sender': sender,
    }
    subject = get_setting(
        'proposal_accepted_subject',
        'email_subject',
        '[abp] Proposal Accepted',
    )
    email.send_email(
        subject,
        context,
        from_email,
        proposal.owner.email,
        email_text,
        proposal=proposal,
        book=submission,
        attachment=attachment,
        request=request,
        kind='proposal',
    )


def send_proposal_revisions(request, proposal, email_text, sender):
    from_email = request.user.email or get_setting('from_address', 'email')
    press_name = get_setting('press_name', 'general')

    if request:
        from_email = "%s <%s>" % (
            request.user.profile.full_name(),
            from_email,
        )

    context = {
        'base_url': get_setting('base_url', 'general'),
        'proposal': proposal,
        'sender': sender,
        'press_name': press_name,
    }
    subject = get_setting(
        'proposal_revision_required_subject',
        'email_subject',
        '[abp] Proposal Revisions Required',
    )
    email.send_email(
        subject,
        context,
        from_email,
        proposal.owner.email,
        email_text,
        proposal=proposal,
        request=request,
        kind='proposal',
    )


def send_proposal_contract_author_sign_off(proposal, email_text, sender):
    from_email = sender.email or get_setting('from_address', 'email')

    context = {
        'base_url': get_setting('base_url', 'general'),
        'proposal': proposal,
        'sender': sender,
    }

    email.send_email(get_setting('book_contract_uploaded_subject',
                                 'email_subject', 'Book Contract Uploaded'),
                     context,
                     from_email,
                     proposal.owner.email,
                     email_text,
                     proposal=proposal,
                     kind='proposal')


def send_invite_typesetter(book, typeset, email_text, sender, attachment):
    from_email = sender.email or get_setting('from_address', 'email')

    context = {
        'base_url': get_setting('base_url', 'general'),
        'submission': typeset.book,
        'typeset': typeset,
        'sender': sender,
    }
    subject = get_setting('typesetting_subject', 'email_subject', 'Typesetting')
    email.send_email(
        subject,
        context,
        from_email,
        typeset.typesetter.email,
        email_text,
        book=book,
        attachment=attachment,
        kind='typeset',
    )


def send_new_user_ack(email_text, new_user, profile):
    from_email = get_setting('from_address', 'email')
    press_name = get_setting('press_name', 'general')

    context = {
        'base_url': get_setting('base_url', 'general'),
        'user': new_user,
        'profile': profile,
        'press_name': press_name,
    }
    subject = get_setting(
        'registration_confirmation_subject',
        'email_subject',
        'Registration Confirmation',
    )
    email.send_email(
        subject,
        context,
        from_email,
        new_user.email,
        email_text,
        kind='general',
    )


def get_active_proposal_form():
    """ Return the current active proposal form.
    Looks for the first form marked as active and
    not in edit (there should only be one). If not,
    returns the first form it can find not in edit.
    """

    active_form = models.ProposalForm.objects.filter(
        active=True,
        in_edit=False
    ).first()

    if not active_form:
        active_form = models.ProposalForm.objects.filter(in_edit=False).first()

    return active_form


def get_file_mimetype(file_path):
    """Returns a guessed mimetype for a given file path.

    Args:
        file_path (str): the path to the storage location of the file.

    Returns:
        str
    """
    mimetype_guess = mimetypes.guess_type(file_path)

    mimetype = mimetype_guess[0]
    if not mimetype:
        mimetype = 'unknown'

    return mimetype


def get_list_of_editors(proposal):
    book_editors = proposal.book_editors.all()
    previous_editors = []
    [previous_editors.append(book_editor) for book_editor in book_editors]
    all_book_editors = User.objects.filter(profile__roles__slug='book-editor')
    list_of_editors = [{} for t in range(0, len(all_book_editors))]

    for t, editor in enumerate(all_book_editors):
        already_added = False
        if editor in previous_editors:
            already_added = True
        list_of_editors[t] = {
            'editor': editor,
            'already_added': already_added,
        }

    return list_of_editors


@is_reviewer
def view_completed_proposal_review(request, proposal_id, assignment_id):
    _proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    proposal_form = manager_forms.GeneratedForm(
        form=models.ProposalForm.objects.get(pk=_proposal.form.id)
    )
    relationships = models.ProposalFormElementsRelationship.objects.filter(
        form=_proposal.form
    )
    data = json.loads(_proposal.data)
    initial_data = {}

    for k, v in data.items():
        initial_data[k] = v[0]

    proposal_form.initial = initial_data
    review_assignment = get_object_or_404(
        submission_models.ProposalReview,
        pk=assignment_id,
        withdrawn=False,
    )
    result = review_assignment.results

    if review_assignment.review_form:
        form = review_forms.GeneratedForm(form=review_assignment.review_form)
    else:
        review_assignment.review_form = _proposal.review_form
        review_assignment.save()
        form = review_forms.GeneratedForm(form=_proposal.review_form)

    ci_required = get_setting('ci_required', 'general')
    recommendation_form = forms.RecommendationForm(
        ci_required=ci_required
    )

    if result:
        relations = review_models.FormElementsRelationship.objects.filter(
            form=result.form
        )
        data_ordered = order_data(
            decode_json(result.data),
            relations
        )
    else:
        data_ordered = None

    if not request.POST and request.GET.get('download') == 'proposal':
        path = create_proposal_form(_proposal)
        return serve_proposal_file(request, path)

    elif not request.POST and request.GET.get('download') == 'docx':
        path = create_completed_proposal_review_form(
            _proposal,
            review_assignment.pk
        )
        return serve_proposal_file(request, path)

    elif request.POST:
        form = review_forms.GeneratedForm(
            request.POST,
            request.FILES,
            form=review_assignment.review_form
        )
        recommendation_form = forms.RecommendationForm(
            request.POST,
            ci_required=ci_required
        )
        if form.is_valid() and recommendation_form.is_valid():
            save_dict = {}
            file_fields = review_models.FormElementsRelationship.objects.filter(
                form=review_assignment.review_form,
                element__field_type='upload'
            )
            data_fields = review_models.FormElementsRelationship.objects.filter(
                ~Q(element__field_type='upload'),
                form=review_assignment.review_form
            )

            for field in file_fields:
                if field.element.name in request.FILES:
                    save_dict[field.element.name] = [
                        review_logic.handle_review_file(
                            request.FILES[field.element.name],
                            'proposal',
                            review_assignment,
                            'reviewer'
                        )
                    ]

            for field in data_fields:
                if field.element.name in request.POST:
                    save_dict[field.element.name] = [
                        request.POST.get(field.element.name),
                        'text'
                    ]

            json_data = smart_text(json.dumps(save_dict))
            form_results = review_models.FormResult(
                form=review_assignment.review_form,
                data=json_data
            )
            form_results.save()

            if request.FILES.get('review_file_upload'):
                review_logic.handle_review_file(
                    request.FILES.get('review_file_upload'), 'proposal',
                    review_assignment, 'reviewer')

            review_assignment.completed = timezone.now()

            if not review_assignment.accepted:
                review_assignment.accepted = timezone.now()
            review_assignment.recommendation = request.POST.get(
                'recommendation'
            )
            review_assignment.competing_interests = request.POST.get(
                'competing_interests'
            )
            review_assignment.results = form_results
            review_assignment.save()

            return redirect(reverse('user_dashboard'))

    template = 'core/proposals/completed_review_assignment.html'
    context = {
        'proposal': _proposal,
        'proposal_form': proposal_form,
        'review_assignment': review_assignment,
        'data_ordered': data_ordered,
        'data_ordered_size': len(data_ordered),
        'result': result,
        'form': form,
        'recommendation_form': recommendation_form,
        'active': 'proposal_review',
        'relationships': relationships,
        'instructions': get_setting(
            'instructions_for_task_proposal',
            'general'
        ),
        'data': data,
    }

    return render(request, template, context)


def create_completed_proposal_review_form(proposal, review_id):
    document = Document()

    if proposal.subtitle:
        document.add_heading("%s: %s" % (proposal.title, proposal.subtitle), 0)
    else:
        document.add_heading(proposal.title, 0)

    review_assignment = get_object_or_404(
        submission_models.ProposalReview,
        pk=review_id,
    )
    if review_assignment.review_form:
        relations = review_models.FormElementsRelationship.objects.filter(
            form=review_assignment.review_form
        ).order_by(
            'order'
        )
    else:
        review_assignment.review_form = proposal.review_form
        review_assignment.save()
        relations = review_models.FormElementsRelationship.objects.filter(
            form=proposal.review_form
        ).order_by(
            'order'
        )

    if review_assignment.results:
        p = document.add_paragraph(
            '%s completed this review assignment form.' %
            review_assignment.user.profile.full_name()
        )

        data = json.loads(review_assignment.results.data)

        for relation in relations:
            v = data.get(relation.element.name)
            document.add_heading(relation.element.name, level=1)
            text = strip_html_tags(smart_text(v[0]))
            document.add_paragraph(text).bold = True
            recommendations = {
                'accept': 'Accept',
                'reject': 'Reject',
                'revisions': 'Revisions Required'
            }

        document.add_heading("Recommendation", level=1)
        document.add_paragraph(
            recommendations[review_assignment.recommendation]
        ).italic = True
        document.add_heading("Competing Interests", level=1)
        document.add_paragraph(
            review_assignment.competing_interests
        ).italic = True

    else:
        p = document.add_paragraph(
            'You should complete this form and then '
            'use the review assignment page to upload it.'
        )

        for relation in relations:

            if (
                    relation.element.field_type in
                    ['text', 'textarea', 'date', 'email']
            ):
                document.add_heading(
                    strip_html_tags(relation.element.name) +
                    ": _______________________________",
                    level=1
                )
                document.add_paragraph(
                    strip_html_tags(relation.help_text)
                ).italic = True

            if relation.element.field_type in ['select', 'check']:
                document.add_heading(
                    strip_html_tags(relation.element.name),
                    level=1
                )

                if relation.element.field_type == 'select':
                    choices = render_choices(relation.element.choices)
                else:
                    choices = ['Y', 'N']

                p = document.add_paragraph(
                    strip_html_tags(relation.help_text)
                )
                p.add_run(
                    ' Mark your choice however you like, '
                    'as long as it is clear.'
                ).italic = True
                table = document.add_table(rows=2, cols=len(choices))
                hdr_cells = table.rows[0].cells

                for i, choice in enumerate(choices):
                    hdr_cells[i].text = choice[0]

                table.style = 'TableGrid'

    document.add_page_break()

    if not os.path.exists(os.path.join(settings.BASE_DIR, 'files', 'forms')):
        os.makedirs(os.path.join(settings.BASE_DIR, 'files', 'forms'))

    path = os.path.join(
        settings.FORM_DIR, '%s.docx' % str(uuid4())
    )

    with default_storage.open(path, 'wb') as file_stream:
        document.save(file_stream)

    return path


def create_proposal_form(proposal):
    document = Document()
    document.add_heading(proposal.title, 0)
    document.add_paragraph(
        'You should complete this form and then '
        'use the proposal page to upload it.'
    )
    relations = models.ProposalFormElementsRelationship.objects.filter(
        form=proposal.form
    ).order_by(
        'order'
    )
    document.add_heading("Title", level=1)
    document.add_paragraph(proposal.title).italic = True
    document.add_heading("Subtitle", level=1)
    document.add_paragraph(proposal.subtitle).italic = True
    document.add_heading("Author", level=1)
    document.add_paragraph(proposal.author).italic = True

    data = json.loads(proposal.data)

    for relation in relations:
        v = data.get(relation.element.name)
        if v:
            document.add_heading(relation.element.name, level=1)
            text = BeautifulSoup(smart_text(v[0]), "html.parser").get_text()
            document.add_paragraph(text).bold = True

    document.add_page_break()

    form_file_path = os.path.join(settings.FORM_DIR, f'{uuid4()}.docx')

    with default_storage.open(form_file_path, 'wb') as file_stream:
        document.save(file_stream)

    return form_file_path


@is_reviewer
def serve_proposal_file(request, file_path):
    try:
        fsock = default_storage.open(file_path, 'r')
        mimetype = get_file_mimetype(file_path)
        response = StreamingHttpResponse(fsock, content_type=mimetype)
        response['Content-Disposition'] = (
            "attachment; filename=proposal_form.docx"
        )
        return response
    except IOError:
        messages.add_message(request, messages.ERROR, 'File not found.')
        return HttpResponseRedirect(request.META.get('HTTP_REFERER'))


def render_choices(choices):
    c_split = choices.split('|')
    return [(choice.capitalize(), choice) for choice in c_split]


@is_reviewer
def create_proposal_review_form(request, proposal):
    document = Document()
    document.add_heading(proposal.proposal.title, 0)
    p = document.add_paragraph(
        'You should complete this form and then '
        'use the review page to upload it.'
    )
    relations = review_models.FormElementsRelationship.objects.filter(
        form=proposal.review_form
    ).order_by(
        'order'
    )

    for relation in relations:

        if relation.element.field_type in ['text', 'textarea', 'date', 'email']:

            document.add_heading(
                strip_html_tags(relation.element.name) +
                ": _______________________________",
                level=1
            )
            document.add_paragraph(
                strip_html_tags(relation.help_text)
            ).italic = True

        if relation.element.field_type in ['select', 'check']:
            document.add_heading(
                strip_html_tags(relation.element.name),
                level=1
            )

            if relation.element.field_type == 'select':
                choices = render_choices(relation.element.choices)
            else:
                choices = ['Y', 'N']

            p = document.add_paragraph(strip_html_tags(relation.help_text))
            p.add_run(
                ' Mark your choice however you like, as long as it is clear.'
            ).italic = True
            table = document.add_table(rows=2, cols=len(choices))
            hdr_cells = table.rows[0].cells

            for i, choice in enumerate(choices):
                hdr_cells[i].text = choice[0]

            table.style = 'TableGrid'

    document.add_page_break()

    path = os.path.join(
        settings.FORM_DIR, '%s.docx' % str(uuid4())
    )

    with default_storage.open(path, 'wb') as file_stream:
        document.save(file_stream)

    return path
