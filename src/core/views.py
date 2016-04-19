from django.contrib.auth import authenticate, logout as logout_user, login as login_user
from django.contrib.auth.models import User
from django.shortcuts import redirect, render, get_object_or_404
from django.http import HttpResponseRedirect, Http404, HttpResponse, StreamingHttpResponse
from django.contrib import messages
from django.core.urlresolvers import reverse
from django.contrib.auth.decorators import login_required
from django import forms
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Q
from django.core import serializers
from django.conf import settings
from django.template import Context, Template
from django.db import IntegrityError
from django.utils.encoding import smart_text

from review import models as review_models
from core import log, models, forms, logic
from author import orcid
from email import send_email,send_reset_email
from files import handle_file_update, handle_attachment, handle_file, handle_proposal_file, handle_proposal_file_form
from submission import models as submission_models
from core.decorators import is_reviewer, is_editor, is_book_editor, is_book_editor_or_author, is_onetasker,is_author, is_press_editor
from review import forms as review_forms
from datetime import datetime
from manager import models as manager_models
from manager import forms as manager_forms
from submission import forms as submission_forms

from docx import Document
from docx.shared import Inches
from pprint import pprint
import json
from time import strftime
import time
from uuid import uuid4
import os
import mimetypes
import mimetypes as mime
from bs4 import BeautifulSoup
import zipfile
import StringIO

from  __builtin__ import any as string_any
import string
# Website Views

def index(request):
    return redirect(reverse('login'))

def contact(request):

    template = "core/contact.html"
    context = {}

    return render(request, template, context)

# Authentication Views
def dashboard(request):
    if request.user.is_authenticated():
        roles=  request.user.profile.roles.all()
        if request.GET.get('next'):
            return redirect(request.GET.get('next'))
        elif string_any('Editor' in role.name for role in roles):
            return redirect(reverse('editor_dashboard'))
        elif string_any('Author' in role.name for role in roles):
            return redirect(reverse('author_dashboard'))
        elif string_any('Reviewer' in role.name for role in roles):
            return redirect(reverse('reviewer_dashboard'))
        else:
            return redirect(reverse('onetasker_dashboard'))
            
def login(request):
    if request.user.is_authenticated():
        messages.info(request, 'You are already logged in.')
        roles=  request.user.profile.roles.all()
        if request.GET.get('next'):
            return redirect(request.GET.get('next'))
        else:
            return redirect(reverse('user_dashboard'))
            

    if request.POST:
        user = request.POST.get('user_name')
        pawd = request.POST.get('user_pass')

        user = authenticate(username=user, password=pawd)

        if user is not None:
            if user.is_active:
                login_user(request, user)
                messages.info(request, 'Login successful.')
                roles=  user.profile.roles.all()
                if request.GET.get('next'):
                    return redirect(request.GET.get('next'))
                else:
                    return redirect(reverse('user_dashboard'))
            else:
                messages.add_message(request, messages.ERROR, 'User account is not active.')
        else:
            messages.add_message(request, messages.ERROR, 'Account not found with those details.')

    context = {}
    template = 'core/login.html'

    return render(request, template, context)

@is_press_editor
def switch_account(request):
    if not request.user.is_authenticated():
       return redirect(reverse('login'))
            
    users = models.Profile.objects.all()
    clean_users = []

    for profile in users:
        user_roles = [role.slug for role in profile.roles.all()]
        if not 'press-editor' in user_roles and not 'book-editor' in user_roles and not 'production-editor' in user_roles:
            if 'author' in user_roles or 'reviewer' in user_roles or 'indexer' in user_roles or 'typesetter' in user_roles or  'copyeditor' in user_roles:
                clean_users.append(profile.user)
    context = {
        'users': clean_users,

    }
    template = 'core/switch_account.html'

    return render(request, template, context)

@is_press_editor
def switch_account_user(request, account_id):
    user = get_object_or_404(User, pk=account_id)
    user.backend = 'django.contrib.auth.backends.ModelBackend'
    login_user(request, user)
    return redirect(reverse('user_dashboard'))

def login_orcid(request):

    orcid_code = request.GET.get('code', None)

    if orcid_code:
        auth = orcid.retrieve_tokens(orcid_code, domain=request.get_host())
        orcid_id = auth.get('orcid', None)

        if orcid_id:
            try:
                user = User.objects.get(profile__orcid=orcid_id)
                user.backend = 'django.contrib.auth.backends.ModelBackend'
                login_user(request, user)

                if request.GET.get('next'):
                    return redirect(request.GET.get('next'))
                else:
                    return redirect(reverse('user_dashboard'))
            except User.DoesNotExist:
                messages.add_message(request, messages.WARNING, 'No user foud with the supplied ORCiD.')
                return redirect(reverse('login'))
        else:
            messages.add_message(request, messages.WARNING, 'Valid ORCiD not returned, please try again, or login with your username and password.')
            return redirect(reverse('login'))
    else:
        messages.add_message(request, messages.WARNING, 'No authorisation code provided, please try again or login with your username and password.')
        return redirect(reverse('login'))

@login_required
def logout(request):
    messages.info(request, 'You have been logged out.')
    logout_user(request)
    return redirect(reverse('index'))

def register(request):
    if request.method == 'POST':
        form = forms.UserCreationForm(request.POST)
        if form.is_valid():
            author_role = models.Role.objects.get(slug='author')
            new_user = form.save()
            new_user.profile.roles.add(author_role)
            new_user.save()

            messages.add_message(request, messages.INFO, models.Setting.objects.get(group__name='general', name='registration_message').value)
            return redirect(reverse('login'))
    else:
        form = forms.UserCreationForm()

    return render(request, "core/register.html", {
        'form': form,
    })

def activate(request, code):
    profile = get_object_or_404(models.Profile, activation_code=code)
    if profile:
        profile.user.is_active = True   
        if not profile.roles.filter(slug='reader').exists():
            profile.roles.add(models.Role.objects.get(slug="reader"))
        if not profile.roles.filter(slug='author').exists():
            profile.roles.add(models.Role.objects.get(slug="author"))
        profile.date_confirmed = timezone.now()
        profile.activation_code = ''
        profile.save()
        profile.user.save()
        messages.add_message(request, messages.INFO, 'Registration complete, you can login now.')
        return redirect(reverse('login'))

@login_required
def view_profile(request):
    try:
        user_profile = models.Profile.objects.get(user=request.user)
    except:
        user_profile = models.Profile(user=request.user)
        user_profile.save()
    name_len=len(request.user.first_name)+len(request.user.last_name)

    template = 'core/user/profile.html'
    context = {
        'user_profile': user_profile,
        'name_width':name_len*14,
        'readonly': False,
        'user_exists': True,
    }

    return render(request, template, context)

@login_required
def view_profile_readonly(request,user_id):
    if request.user.pk == user_id:
        return redirect(reverse('view_profile'))
    user_exists = False
    user_profile = None
    try:
        user_profile = models.Profile.objects.get(user__pk=user_id)
        user_exists = True
    except:
        user_exists = False

    name_len=len(request.user.first_name)+len(request.user.last_name)

    template = 'core/user/profile.html'
    context = {
        'user_profile': user_profile,
        'name_width':name_len*14,
        'readonly': True,
        'user_exists': user_exists,
    }

    return render(request, template, context)

@login_required
def update_profile(request):
    user_profile = models.Profile.objects.get(user=request.user)
    user_form = forms.UserProfileForm(instance=request.user)
    profile_form = forms.ProfileForm(instance=user_profile)
    if request.method == 'POST':
        user_form = forms.UserProfileForm(request.POST, instance=request.user)
        profile_form = forms.ProfileForm(request.POST, request.FILES, instance=user_profile)
        if profile_form.is_valid() and user_form.is_valid():
            user = user_form.save()
            profile = profile_form.save()
            for interest in profile.interest.all():
                profile.interest.remove(interest)
            
            interests = request.POST.get('interests')
            if interests:
                for interest in interests.split(','):
                    new_interest, c = models.Interest.objects.get_or_create(name=interest)
                    profile.interest.add(new_interest)
            profile.save()

            return redirect(reverse('view_profile'))

    template = 'core/user/update_profile.html'
    context = {
        'profile_form' : profile_form,
        'user_form': user_form,
        'user':request.user,
    }

    return render(request, template, context)

def oai(request):
    
    base_url = models.Setting.objects.get(group__name='general', name='base_url').value
    oai_dc = 'http://%s/oai' % (base_url)
    oai_identifier = models.Setting.objects.get(group__name='general', name='oai_identifier').value

    books = models.Book.objects.all()
    list_of_books =[[{}] for t in range(0,len(books)) ] 

    for t,book in enumerate(books):
        try:
            isbns = models.Identifier.objects.filter(book=book).exclude(identifier='pub_id').exclude(identifier='urn').exclude(identifier='doi')
        except:
            isbns=None
        formats=book.formats()
        list_format=[]
        for format in formats:
            list_format.append(format.file.mime_type)
        list_of_books[t] = [{'book': book, 'isbns': isbns,'formats':list_format,}]

    template = 'core/oai.xml'
    context = {
        'books': list_of_books,
        'oai_dc': oai_dc,
        'base_url':base_url,
        'oai_identifier': oai_identifier,
    }

    return render(request, template, context,content_type="application/xhtml+xml")

@login_required
def user_home(request):

    task_list = models.Task.objects.filter(assignee=request.user, completed__isnull=True).order_by('due')
    new_task_form = forms.TaskForm()

    template = 'core/user/home.html'
    context = {
        'task_list': task_list,
        'proposals': submission_models.Proposal.objects.filter(status='submission').count(),
        'new_submissions': models.Book.objects.filter(stage__current_stage='submission').count(),
        'in_review': models.Book.objects.filter(stage__current_stage='review').count(),
        'in_editing': models.Book.objects.filter(stage__current_stage='editing').count(),
        'in_production': models.Book.objects.filter(stage__current_stage='production').count(),
        'copyedits': models.CopyeditAssignment.objects.filter(copyeditor=request.user, completed__isnull=True),
        'new_task_form': new_task_form,
        'user_submissions': models.Book.objects.filter(owner=request.user),
        'author_copyedit_tasks': logic.author_tasks(request.user),
        'indexes': models.IndexAssignment.objects.filter(indexer=request.user, completed__isnull=True),
        'typesetting': models.TypesetAssignment.objects.filter((Q(requested__isnull=False) & Q(completed__isnull=True)) | (Q(typesetter_invited__isnull=False) & Q(typesetter_completed__isnull=True)), typesetter=request.user),
        'user_proposals': submission_models.Proposal.objects.filter(owner=request.user)
    }

    return render(request, template, context)

@login_required
def user_submission(request, submission_id):
    book = get_object_or_404(models.Book, pk=submission_id, owner=request.user)

    template = 'core/user/user_submission.html'
    context = {
        'submission': book,
        'active': 'user_submission',
    }

    return render(request, template, context)

@login_required
def reset_password(request):

    if request.method == 'POST':
        password_1 = request.POST.get('password_1')
        password_2 = request.POST.get('password_2')

        if password_1 == password_2:
            if len(password_1) > 8:
                user = User.objects.get(username=request.user.username)
                user.set_password(password_1)
                user.save()
                messages.add_message(request, messages.SUCCESS, 'Password successfully changed.')
                return redirect(reverse('login'))
            else:
                messages.add_message(request, messages.ERROR, 'Password is not long enough, must be greater than 8 characters.')
        else:
            messages.add_message(request, messages.ERROR, 'Your passwords do not match.')

    template = 'core/user/reset_password.html'
    context = {}

    return render(request, template, context)

def unauth_reset(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        try:
            user = User.objects.get(username=username)          
            password = uuid4()
            user.profile.reset_code=password
            user.profile.save()
            user.save()
            email_text = models.Setting.objects.get(group__name='email', name='reset_password').value
            send_reset_email(user=user, email_text=email_text, reset_code=password)
            messages.add_message(request, messages.INFO, 'A reset code has been sent to your email account.')
        except User.DoesNotExist:
            messages.add_message(request, messages.ERROR, 'There is no account for that username.')

    template = 'core/user/reset_username.html'
    context = {}

    return render(request, template, context)

def unauth_reset_code(request, uuid):
    valid_user = False  
    try:
        user = get_object_or_404(User,profile__reset_code=uuid)
        if user.profile.reset_code:
            valid_user = True
        user.profile.reset_code_validated = True
        user.profile.save()
        user.save()
        messages.add_message(request, messages.SUCCESS, 'Valid reset code')

        return redirect(reverse('unauth_reset_password',kwargs={'uuid':uuid}))
            
    except User.DoesNotExist:
        messages.add_message(request, messages.ERROR, 'There is no account associated with that reset code.')
        return redirect(reverse('login'))
  
def unauth_reset_password(request, uuid):
    try:
        user = get_object_or_404(User,profile__reset_code=uuid)        
    except User.DoesNotExist:
        user = None
        messages.add_message(request, messages.ERROR, 'There is no account for that username or reset code is invalid.')
    if user:
        valid_reset = user.profile.reset_code_validated
    else:
        valid_reset = False
    if valid_reset:
        if request.method == 'POST':
            password_1 = request.POST.get('password_1')
            password_2 = request.POST.get('password_2')

            if password_1 == password_2:
                if len(password_1) > 8:
                    user = User.objects.get(username=user.username)
                    user.set_password(password_1)
                    user.save()
                    user.profile.reset_code_validated = False
                    user.profile.reset_code = None
                    user.profile.save()
                    user.save()
                    messages.add_message(request, messages.SUCCESS, 'Password successfully changed.')
                    return redirect(reverse('login'))
                else:
                    messages.add_message(request, messages.ERROR, 'Password is not long enough, must be greater than 8 characters.')
            else:
                messages.add_message(request, messages.ERROR, 'Your passwords do not match.')

    template = 'core/user/reset_password_unauth.html'
    context = {
    'valid_reset':valid_reset,
    'user':user,

    }

    return render(request, template, context)

def permission_denied(request):

    template = 'core/403.html'
    context = {}

    return render(request, template, context)

# Dashboard
@login_required
def overview(request):

    template = 'core/dashboard/dashboard.html'
    context = {
        'new_submissions': models.Book.objects.filter(stage__current_stage='submission'),
        'in_review': models.Book.objects.filter(stage__current_stage='review'),
        'in_editing': models.Book.objects.filter(stage__current_stage='editing'),
        'in_production': models.Book.objects.filter(stage__current_stage='production'),
    }

    return render(request, template, context)
@login_required
def overview_inprogress(request):

    template = 'core/dashboard/inprogress_dashboard.html'
    context = {
        'submissions': models.Book.objects.filter(stage__isnull=True),
    }

    return render(request, template, context)
@login_required
def proposal_overview(request):

    template = 'core/dashboard/proposal_overview.html'
    context = {
        'proposals': submission_models.Proposal.objects.exclude(status='declined').exclude(status='accepted'),
        'declined_proposals': submission_models.Proposal.objects.filter(status='declined'),
        'accepted_proposals': submission_models.Proposal.objects.filter(status='accepted'),  
    }

    return render(request, template, context)

# AJAX Handlers
@csrf_exempt
@login_required
def task_complete(request, task_id):

    task = get_object_or_404(models.Task, pk=task_id, assignee=request.user, completed__isnull=True)
    task.completed = timezone.now()
    task.save()
    return HttpResponse('Thanks')

@login_required
def task_new(request):

    new_task_form = forms.TaskForm(request.POST)
    if new_task_form.is_valid():
        task = new_task_form.save(commit=False)
        task.creator = request.user
        task.assignee = request.user
        task.save()

        return HttpResponse(smart_text(json.dumps({'id': task.pk,'text': task.text})))
    else:
        return HttpResponse(new_task_form.errors)

@csrf_exempt
@is_book_editor_or_author
def new_message(request, submission_id):

    new_message_form = forms.MessageForm(request.POST)
    if new_message_form.is_valid():
        new_message = new_message_form.save(commit=False)
        new_message.sender = request.user
        new_message.book = get_object_or_404(models.Book, pk=submission_id)
        new_message.save()

        response_dict = {
            'status_code': 200, 
            'message_id': new_message.pk,
            'sender': new_message.sender.profile.full_name(),
            'message': new_message.message,
            'date_sent': new_message.date_sent.strftime("%-d %b %Y, %H:%M"),
        }

        return HttpResponse(smart_text(json.dumps(response_dict)))
    else:
        return HttpResponse(smart_text(json.dumps(new_message_form.errors)))

@csrf_exempt
@is_book_editor_or_author
def get_messages(request, submission_id):
    try:
        last_message = int(request.GET.get('last_message', 0))
        messages = models.Message.objects.filter(book__pk=submission_id, pk__gt=last_message).exclude(sender=request.user).order_by('-id')

        message_list = [{
                'message': message.message,
                'message_id': message.pk,
                'sender': message.sender.profile.full_name(),
                'initials': message.sender.profile.initials(),
                'message': message.message,
                'date_sent': message.date_sent.strftime("%-d %b %Y, %H:%M"),
                'user':'same',
            } for message in messages
        ]
        response_dict = {
            'status_code': 200,
            'messages': message_list,
        }

        return HttpResponse(smart_text(json.dumps(response_dict)))
    except Exception, e:
        return HttpResponse(e)

def get_authors(request, submission_id):
    if request.is_ajax():
        q = request.GET.get('term', '')
        data = smart_text(json.dumps(logic.get_author_emails(submission_id,q)))
    else:
        data = 'Unable to get authors'
    mimetype = 'application/json'
    return HttpResponse(data, mimetype)

def get_editors(request, submission_id):
    if request.is_ajax():
        q = request.GET.get('term', '')
        results = []
        data = smart_text(json.dumps(logic.get_editor_emails(submission_id,q)))
    else:
        data = 'Unable to get editors'

    mimetype = 'application/json'
    return HttpResponse(data, mimetype)

def get_onetaskers(request, submission_id):
    if request.is_ajax():
        q = request.GET.get('term', '')
        data = smart_text(json.dumps(logic.get_onetasker_emails(submission_id,q)))
    else:
        data = 'Unable to get onetaskers'
    mimetype = 'application/json'
    return HttpResponse(data, mimetype)

def get_all(request, submission_id):
    submission = get_object_or_404(models.Book, pk=submission_id)
    if request.is_ajax():
        q = request.GET.get('term', '')
        onetasker_results = logic.get_onetasker_emails(submission_id,q)
        editor_results = logic.get_editor_emails(submission_id,q)
        author_results = logic.get_author_emails(submission_id,q)
        results = []
        for user in onetasker_results:
            if not string_any(user['value'] in result['value'] for result in results):

               results.append(user)
        for author in author_results:
            if not string_any(author['value'] in result['value'] for result in results):
                results.append(author)
                
        for editor in editor_results:
            if not string_any(editor['value'] in result['value'] for result in results):
                results.append(editor)
        data = smart_text(json.dumps(results))
    else:
        data = 'Unable to get any user'
    mimetype = 'application/json'
    return HttpResponse(data, mimetype)

def get_proposal_users(request, proposal_id):
    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    if request.is_ajax():
        q = request.GET.get('term', '')
        proposal_results = logic.get_proposal_emails(proposal_id,q)
        results = []
        for user in proposal_results:
            if not string_any(user['value'] in result['value'] for result in results):
               results.append(user)
        data = smart_text(json.dumps(results))
    else:
        data = 'Unable to get any user'
    mimetype = 'application/json'
    return HttpResponse(data, mimetype)

@login_required
def email_users(request, group, submission_id=None, user_id=None):
    submission = get_object_or_404(models.Book, pk=submission_id)
    editors = logic.get_editors(submission)
    authors = submission.author.all()
    onetaskers = submission.onetaskers()
    to_value=""
    sent = False
    if request.POST:

        attachment = request.FILES.get('attachment')
        subject = request.POST.get('subject')
        body = request.POST.get('body')
        
        to_addresses = request.POST.get('to_values').split(';')
        cc_addresses = request.POST.get('cc_values').split(';')
        bcc_addresses = request.POST.get('bcc_values').split(';')

        to_list=logic.clean_email_list(to_addresses)
        cc_list=logic.clean_email_list(cc_addresses)
        bcc_list=logic.clean_email_list(bcc_addresses)
        
        if attachment: 
            attachment = handle_file(attachment, submission, 'other', request.user, "Attachment: Uploaded by %s" % (request.user.username))
        
        if to_addresses:
            if attachment: 
                send_email(subject=subject, context={}, from_email=request.user.email, to=to_list, bcc=bcc_list,cc=cc_list, html_template=body, book=submission, attachment=attachment)
            else:
                send_email(subject=subject, context={}, from_email=request.user.email, to=to_list,bcc=bcc_list,cc=cc_list, html_template=body, book=submission)
            message ="E-mail with subject '%s' was sent." % (subject)
            return HttpResponse('<script type="text/javascript">window.alert("'+message+'")</script><script type="text/javascript">window.close()</script>') 

    if not group == "all" and user_id:
        
        if group == "editors":
            try:
                editor = models.User.objects.get(pk=user_id)
                if editor in editors:
                    to_value="%s;" % (editor.email)
                else:
                    messages.add_message(request, messages.ERROR, "This editor is not an editor of this submission")
            except models.User.DoesNotExist:
                messages.add_message(request, messages.ERROR, "This editor was not found")

        elif group == "authors":
            author = get_object_or_404(models.Author, pk=user_id)
            authors = submission.author.all()
            if author in authors:
                to_value="%s;" % (author.author_email)
            else:
                messages.add_message(request, messages.ERROR, "This author is not an author of this submission")

        elif group == "onetaskers":
            user = get_object_or_404(models.User, pk=user_id)
            if user in onetaskers:
                to_value="%s;" % (user.email)
            else:
                messages.add_message(request, messages.ERROR, "This onetasker was not found")

    elif group =="all" and user_id:
        messages.add_message(request, messages.ERROR, "Cannot use the user field on this page because of the 'all' in the url. Try replacing it with other email groups: 'authors' or 'editors' or 'onetaskers'")
    
    group_name=group
    
    if not group_name == "editors" and not group_name == "all" and not group_name == "authors" and not group_name == "onetaskers":
        messages.add_message(request, messages.ERROR, "Group type does not exist. Redirected to page of all groups")
        return redirect(reverse('email_users', kwargs={'group':'all','submission_id': submission.id}))
    
    source = "/email/get/%s/submission/%s/" % (group_name,submission_id)

    template = 'core/email.html'
    context = {
        'submission': submission,
        'from': request.user,
        'to_value':to_value,
        'source': source,
        'group': group_name,
        'user_id':user_id,
        'sent':sent,
        
    }
    return render(request, template, context)

@login_required
def email_users_proposal(request, proposal_id, user_id):
    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    users = User.objects.all()
    user = User.objects.get(pk = user_id)
    to_value=""
    sent = False
    if request.POST:

        attachment = request.FILES.get('attachment')
        subject = request.POST.get('subject')
        body = request.POST.get('body')
        
        to_addresses = request.POST.get('to_values').split(';')
        cc_addresses = request.POST.get('cc_values').split(';')
        bcc_addresses = request.POST.get('bcc_values').split(';')

        to_list=logic.clean_email_list(to_addresses)
        cc_list=logic.clean_email_list(cc_addresses)
        bcc_list=logic.clean_email_list(bcc_addresses)
        
        if attachment: 
            attachment = handle_proposal_file(attachment, proposal, 'other', request.user, "Attachment: Uploaded by %s" % (request.user.username))
        
        if to_addresses:
            if attachment: 
                send_email(subject=subject, context={}, from_email=request.user.email, to=to_list, bcc=bcc_list,cc=cc_list, html_template=body, proposal=proposal, attachment=attachment)
            else:
                send_email(subject=subject, context={}, from_email=request.user.email, to=to_list,bcc=bcc_list,cc=cc_list, html_template=body, proposal=proposal)
            message ="E-mail with subject '%s' was sent." % (subject)
            return HttpResponse('<script type="text/javascript">window.alert("'+message+'")</script><script type="text/javascript">window.close()</script>') 

    if not proposal.owner == user and not proposal.requestor == user :
        messages.add_message(request, messages.ERROR, "This user is not associated with this proposal")
    else:
        to_value="%s;" % (user.email)
    source = "/email/user/proposal/%s/" % proposal_id


    template = 'core/email.html'
    context = {
        'proposal': proposal,
        'from': request.user,
        'to_value':to_value,
        'source': source,
        'group': 'proposal',
        'user_id':user_id,
        'sent':sent,
        
    }
    return render(request, template, context)

def page(request, page_name):

    page_content = get_object_or_404(models.Setting, group__name='page', name=page_name)
    title = page_name.replace('-', ' ')

    template = 'core/page.html'
    context = {
        'page_content': page_content,
        'title':title.replace('_', ' ')
    }

    return render(request, template, context)

@is_book_editor
def upload_misc_file(request, submission_id):

    submission = get_object_or_404(models.Book, pk=submission_id)
    if request.POST:
        file_form = forms.UploadMiscFile(request.POST)
        if file_form.is_valid():
            new_file = handle_file(request.FILES.get('misc_file'), submission, file_form.cleaned_data.get('file_type'), request.user, file_form.cleaned_data.get('label'))
            submission.misc_files.add(new_file)
            return redirect(reverse('editor_submission', kwargs={'submission_id': submission.id}))
    else:
        file_form = forms.UploadMiscFile()

    template = 'core/misc_files/upload.html'
    context = {
        'submission': submission,
        'file_form': file_form,
        'active_page':'editor_submission'
    }

    return render(request, template, context)

@is_book_editor
def upload_manuscript(request, submission_id):

    submission = get_object_or_404(models.Book, pk=submission_id)
    if request.POST:
        file_form = forms.UploadFile(request.POST)
        if file_form.is_valid():
            new_file = handle_file(request.FILES.get('manuscript'), submission, 'manuscript', request.user, file_form.cleaned_data.get('label'))
            submission.files.add(new_file)
            return redirect(reverse('editor_submission', kwargs={'submission_id': submission.id}))
    else:
        file_form = forms.UploadFile()

    template = 'core/misc_files/upload_manuscript.html'
    context = {
        'submission': submission,
        'file_form': file_form,
        'active_page':'editor_submission'
    }

    return render(request, template, context)

@is_book_editor
def upload_additional(request, submission_id):

    submission = get_object_or_404(models.Book, pk=submission_id)
    if request.POST:
        file_form = forms.UploadFile(request.POST)
        if file_form.is_valid():
            new_file = handle_file(request.FILES.get('additional'), submission, 'additional', request.user, file_form.cleaned_data.get('label'))
            submission.files.add(new_file)
            return redirect(reverse('editor_submission', kwargs={'submission_id': submission.id}))
    else:
        file_form = forms.UploadFile()

    template = 'core/misc_files/upload_additional.html'
    context = {
        'submission': submission,
        'file_form': file_form,
        'active_page':'editor_submission'
    }

    return render(request, template, context)
    
@login_required
def serve_marc21_file(request, submission_id,type):
    book = get_object_or_404(models.Book, pk=submission_id)
    if type=='xml':
        file_pk=logic.book_to_mark21_file(book,request.user,True)
    else:
        file_pk=logic.book_to_mark21_file(book,request.user)
    _file = get_object_or_404(models.File, pk=file_pk)
    file_path = os.path.join(settings.BOOK_DIR, submission_id, _file.uuid_filename)

    try:
        fsock = open(file_path, 'r')
        mimetype = mimetypes.guess_type(file_path)
        response = StreamingHttpResponse(fsock, content_type=mimetype)
        response['Content-Disposition'] = "attachment; filename=%s" % (_file.original_filename)
        #log.add_log_entry(book=book, user=request.user, kind='file', message='File %s downloaded.' % _file.uuid_filename, short_name='Download')
        return response
    except IOError:
        messages.add_message(request, messages.ERROR, 'File not found. %s' % (file_path))
        return HttpResponseRedirect(request.META.get('HTTP_REFERER'))
        
#reference: http://stackoverflow.com/questions/12881294/django-create-a-zip-of-multiple-files-and-make-it-downloadable
@is_book_editor_or_author
def serve_all_files(request, submission_id):
    book = get_object_or_404(models.Book, pk=submission_id)
    files = book.files.all()
    misc_files = book.misc_files.all()
    internal_review_files = book.internal_review_files.all()
    external_review_files = book.external_review_files.all()
    
    zip_subdir = "Files of Submission #%s" % submission_id
    zip_filename = "%s.zip" % zip_subdir

    # Open StringIO to grab in-memory ZIP contents
    s = StringIO.StringIO()

    # The zip compressor
    zf = zipfile.ZipFile(s, "w")

    for file in files:
        fpath = os.path.join(settings.BOOK_DIR, submission_id, file.uuid_filename)
        fdir, fname = os.path.split(fpath)
        zip_path = os.path.join(zip_subdir, fname)
        zf.write(fpath, zip_path)

    for file in external_review_files:
        fpath = os.path.join(settings.BOOK_DIR, submission_id, file.uuid_filename)
        fdir, fname = os.path.split(fpath)
        zip_path = os.path.join(zip_subdir, fname)
        zf.write(fpath, zip_path)

    for file in internal_review_files:
        fpath = os.path.join(settings.BOOK_DIR, submission_id, file.uuid_filename)
        fdir, fname = os.path.split(fpath)
        zip_path = os.path.join(zip_subdir, fname)
        zf.write(fpath, zip_path)

    for file in misc_files:
        fpath = os.path.join(settings.BOOK_DIR, submission_id, file.uuid_filename)
        fdir, fname = os.path.split(fpath)
        zip_path = os.path.join(zip_subdir, fname)
        zf.write(fpath, zip_path)
    zf.close()
    fd = open(os.path.join(settings.BOOK_DIR, submission_id, zip_filename), 'wb')
    fd.write(s.getvalue())
    fd.close()
    fsock = open(os.path.join(settings.BOOK_DIR, submission_id, zip_filename), 'r')
    resp = StreamingHttpResponse(fsock, content_type = "application/x-zip-compressed")
    os.remove(os.path.join(settings.BOOK_DIR, submission_id, zip_filename))
    resp['Content-Disposition'] = 'attachment; filename=%s' % zip_filename
    return resp


@is_onetasker
def serve_all_review_files(request, submission_id,review_type):
    book = get_object_or_404(models.Book, pk=submission_id)
    internal_review_files = book.internal_review_files.all()
    external_review_files = book.external_review_files.all()
    
    zip_subdir = "Files of Submission #%s" % submission_id
    zip_filename = "%s.zip" % zip_subdir

    # Open StringIO to grab in-memory ZIP contents
    s = StringIO.StringIO()

    # The zip compressor
    zf = zipfile.ZipFile(s, "w")

    if review_type == "external":
        for file in external_review_files:
            fpath = os.path.join(settings.BOOK_DIR, submission_id, file.uuid_filename)
            fdir, fname = os.path.split(fpath)
            zip_path = os.path.join(zip_subdir, fname)
            zf.write(fpath, zip_path)

    else:
        for file in internal_review_files:
            fpath = os.path.join(settings.BOOK_DIR, submission_id, file.uuid_filename)
            fdir, fname = os.path.split(fpath)
            zip_path = os.path.join(zip_subdir, fname)
            zf.write(fpath, zip_path)

    zf.close()
    fd = open(os.path.join(settings.BOOK_DIR, submission_id, zip_filename), 'wb')
    fd.write(s.getvalue())
    fd.close()
    fsock = open(os.path.join(settings.BOOK_DIR, submission_id, zip_filename), 'r')
    resp = StreamingHttpResponse(fsock, content_type = "application/x-zip-compressed")
    os.remove(os.path.join(settings.BOOK_DIR, submission_id, zip_filename))
    resp['Content-Disposition'] = 'attachment; filename=%s' % zip_filename
    return resp

@is_onetasker
def serve_file(request, submission_id, file_id):
    book = get_object_or_404(models.Book, pk=submission_id)
    _file = get_object_or_404(models.File, pk=file_id)
    file_path = os.path.join(settings.BOOK_DIR, submission_id, _file.uuid_filename)

    try:
        fsock = open(file_path, 'r')
        mimetype = mimetypes.guess_type(file_path)
        response = StreamingHttpResponse(fsock, content_type=mimetype)
        response['Content-Disposition'] = "attachment; filename=%s" % (_file.original_filename)
        #log.add_log_entry(book=book, user=request.user, kind='file', message='File %s downloaded.' % _file.original_filename, short_name='Download')
        return response
    except IOError:
        messages.add_message(request, messages.ERROR, 'File not found. %s' % (file_path))
        return HttpResponseRedirect(request.META.get('HTTP_REFERER'))

@is_editor
def serve_proposal_file_id(request, proposal_id, file_id):
    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    _file = get_object_or_404(models.File, pk=file_id)
    file_path = os.path.join(settings.BASE_DIR, 'files', 'proposals',str(proposal_id), _file.uuid_filename)
    try:
        fsock = open(file_path, 'r')
        mimetype = mimetypes.guess_type(file_path)
        response = StreamingHttpResponse(fsock, content_type=mimetype)
        response['Content-Disposition'] = "attachment; filename=%s" % (_file.original_filename)
        return response
    except IOError:
        messages.add_message(request, messages.ERROR, 'File not found. %s' % (file_path))
        return HttpResponseRedirect(request.META.get('HTTP_REFERER'))

@is_editor
def serve_proposal_file(request, proposal_id, file_id):
    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    _file = get_object_or_404(models.File, pk=file_id)
    file_path = os.path.join(settings.PROPOSAL_DIR, proposal_id, _file.original_filename)

    try:
        fsock = open(file_path, 'r')
        mimetype = mimetypes.guess_type(file_path)
        response = StreamingHttpResponse(fsock, content_type=mimetype)
        response['Content-Disposition'] = "attachment; filename=%s" % (_file.original_filename)
        #log.add_proposal_log_entry(proposal=proposal, user=request.user, kind='file', message='File %s downloaded.' % _file.original_filename, short_name='Download')
        return response
    except IOError:
        messages.add_message(request, messages.ERROR, 'File not found. %s' % (file_path))
        return HttpResponseRedirect(request.META.get('HTTP_REFERER'))

@is_book_editor_or_author
def serve_versioned_file(request, submission_id, revision_id):
    book = get_object_or_404(models.Book, pk=submission_id)
    versions_file = get_object_or_404(models.FileVersion, pk=revision_id)
    file_path = os.path.join(settings.BOOK_DIR, submission_id, versions_file.original_filename)

    try:
        fsock = open(file_path, 'r')
        mimetype = mimetypes.guess_type(file_path)
        response = StreamingHttpResponse(fsock, content_type=mimetype)
        response['Content-Disposition'] = "attachment; filename=%s" % (versions_file.original_filename)
        #log.add_log_entry(book=book, user=request.user, kind='file', message='File %s downloaded.' % versions_file.uuid_filename, short_name='Download')
        return response
    except IOError:
        messages.add_message(request, messages.ERROR, 'File not found. %s/%s' % (file_path, versions_file.uuid_filename))
        return HttpResponseRedirect(request.META.get('HTTP_REFERER'))

@is_book_editor_or_author
def delete_file(request, submission_id, file_id, returner):
    book = get_object_or_404(models.Book, pk=submission_id)
    _file = get_object_or_404(models.File, pk=file_id)
    file_id = _file.id
    _file.delete()

    if returner == 'new':
        return redirect(reverse('editor_submission', kwargs={'submission_id': book.id}))
    elif returner == 'review':
        return redirect(reverse('editor_review', kwargs={'submission_id': book.id}))
    elif returner == 'production':
        return redirect(reverse('editor_production', kwargs={'submission_id': book.id}))
    elif returner == 'editor':
        return redirect(reverse('editor_submission', kwargs={'submission_id': book.id}))
    elif returner == 'author':
        return redirect(reverse('author_dashboard'))


@is_book_editor_or_author
def update_file(request, submission_id, file_id, returner):
    book = get_object_or_404(models.Book, pk=submission_id)
    _file = get_object_or_404(models.File, pk=file_id)
    if request.POST:
        label = request.POST['rename']
        if label:
            _file.label=label
            _file.save()

        for file in request.FILES.getlist('update_file'):
            handle_file_update(file, _file, book, _file.kind, request.user)
            messages.add_message(request, messages.INFO, 'File updated.')

        if returner == 'new':
            return redirect(reverse('editor_submission', kwargs={'submission_id': book.id}))
        elif returner == 'review':
            return redirect(reverse('editor_review', kwargs={'submission_id': book.id}))
        elif returner == 'production':
            return redirect(reverse('editor_production', kwargs={'submission_id': book.id}))
        elif returner == 'editor':
            return redirect(reverse('editor_submission', kwargs={'submission_id': book.id}))

    template = 'core/update_file.html'
    context = {
        'submission': book,
        'file': _file,
        'update': True
    }

    return render(request, template, context)

@is_book_editor_or_author
def view_file(request, submission_id, file_id):
    book = get_object_or_404(models.Book, pk=submission_id)
    _file = get_object_or_404(models.File, pk=file_id)
    
    template = 'core/update_file.html'
    context = {
        'submission': book,
        'file': _file,
        'update': False
    }

    return render(request, template, context)

@is_book_editor_or_author
def versions_file(request, submission_id, file_id):
    book = get_object_or_404(models.Book, pk=submission_id)
    _file = get_object_or_404(models.File, pk=file_id)
    versions = models.FileVersion.objects.filter(file=_file).extra(order_by = ['date_uploaded'])

    template = 'core/versions_file.html'
    context = {
        'submission': book,
        'file': _file,
        'versions': versions
    }

    return render(request, template, context)

# Log
@is_book_editor
def view_log(request, submission_id):
    book = get_object_or_404(models.Book, pk=submission_id)
    log_list = models.Log.objects.filter(Q(book=book) | Q(proposal=book.proposal)).order_by('-date_logged')
    email_list = models.EmailLog.objects.filter(Q(book=book) | Q(proposal=book.proposal)).order_by('-sent')

    template = 'editor/log.html'
    context = {
        'submission': book,
        'log_list': log_list,
        'email_list': email_list,
        'active': 'log',
    }

    return render(request, template, context)

## PROPOSALS ##
@is_book_editor
def view_proposal_log(request, proposal_id):
    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    log_list = models.Log.objects.filter(proposal=proposal).order_by('-date_logged')
    email_list = models.EmailLog.objects.filter(proposal=proposal).order_by('-sent')

    template = 'editor/proposal_log.html'
    context = {
        'proposal': proposal,
        'log_list': log_list,
        'email_list': email_list,
        'active': 'log',
    }

    return render(request, template, context)


@is_editor
def assign_proposal(request):
    proposal_form_id = models.Setting.objects.get(name='proposal_form').value
    proposal_form = manager_forms.GeneratedForm(form=models.ProposalForm.objects.get(pk=proposal_form_id))
    default_fields = manager_forms.DefaultForm()

    if request.method == 'POST':
        proposal_form = manager_forms.GeneratedForm(request.POST, request.FILES,form=models.ProposalForm.objects.get(pk=proposal_form_id))
        default_fields = manager_forms.DefaultForm(request.POST)
        if proposal_form.is_valid() and default_fields.is_valid():
            defaults = {field.name: field.value() for field in default_fields}
            proposal = submission_models.Proposal(form=models.ProposalForm.objects.get(pk=proposal_form_id), data=None, owner=None, **defaults)
            proposal.save()
            save_dict = {}
            file_fields = models.ProposalFormElementsRelationship.objects.filter(form=models.ProposalForm.objects.get(pk=proposal_form_id), element__field_type='upload')
            data_fields = models.ProposalFormElementsRelationship.objects.filter(~Q(element__field_type='upload'), form=models.ProposalForm.objects.get(pk=proposal_form_id))

            for field in file_fields:
                if field.element.name in request.FILES:
                    # TODO change value from string to list [value, value_type]
                    save_dict[field.element.name] = [handle_proposal_file_form(request.FILES[field.element.name], proposal, 'other', request.user, "Attachment: Uploaded by %s" % (request.user.username))]

            for field in data_fields:
                if field.element.name in request.POST:
                    # TODO change value from string to list [value, value_type]
                    save_dict[field.element.name] = [request.POST.get(field.element.name), 'text']
            
            json_data = smart_text(json.dumps(save_dict))
            proposal.data = json_data
            proposal.save()
            editors = User.objects.filter(profile__roles__slug='press-editor')
            message = "A new Unassigned Proposal '%s' with id %s has been submitted by %s ."  % (proposal.title,proposal.pk,request.user.username)
            for editor in editors:
                notification = models.Task(assignee=editor,creator=request.user,text=message,workflow='proposal')
                notification.save()

            messages.add_message(request, messages.SUCCESS, 'Unassigned Proposal %s submitted' % proposal.id)
 #           email_text = models.Setting.objects.get(group__name='email', name='proposal_submission_ack').value
  #          logic.send_proposal_submission_ack(proposal, email_text=email_text, owner=request.user)

            log.add_proposal_log_entry(proposal=proposal,user=request.user, kind='proposal', message='Unassigned Proposal has been submitted by %s.' % request.user.profile.full_name(), short_name='Unassigned Proposal Submitted')
    
            return redirect(reverse('proposals',kwargs = {}))


    template = "core/proposals/assign/start_proposal.html"
    context = {
        'proposal_form': proposal_form,
        'default_fields': default_fields,
        'core_proposal':models.ProposalForm.objects.get(pk=proposal_form_id),
    }

    return render(request, template, context)

@is_editor
def proposal_assign_user(request, proposal_id,user_id):

    proposal = submission_models.Proposal.objects.get(pk=proposal_id)
    user = models.User.objects.get(pk=user_id)
    proposal.owner = user 
    proposal.save()
    email_text = models.Setting.objects.get(group__name='email', name='proposal_submission_ack').value
    logic.send_proposal_submission_ack(proposal, email_text=email_text, owner=user)
    messages.add_message(request, messages.SUCCESS, 'Unassigned Proposal %s assigned' % proposal.id)
    log.add_proposal_log_entry(proposal=proposal,user=request.user, kind='proposal', message='Proposal "%s %s" assigned to %s %s.'%(proposal.title,proposal.subtitle,user.first_name,user.last_name), short_name='Proposal Assigned')
          
    return redirect(reverse('proposals'))


@is_editor
def proposal_assign_view(request, proposal_id):

    proposal = submission_models.Proposal.objects.get(pk=proposal_id)
    proposal_form_id = models.Setting.objects.get(name='proposal_form').value
    authors = User.objects.filter(profile__roles__slug='author')
                
    if proposal.owner == request.user:
        viewable = True

    proposal_form = manager_forms.GeneratedForm(form=models.ProposalForm.objects.get(pk=proposal.form.id))
    default_fields = manager_forms.DefaultForm(initial={'title': proposal.title,'author':proposal.author,'subtitle':proposal.subtitle})

    intial_data={}
    data = {}
    if proposal.data:
        data = json.loads(proposal.data)
        for k,v in data.items():
            intial_data[k] = v[0]


    proposal_form.initial=intial_data
    
    roles = request.user.profile.roles.all()

    user_roles = [role.slug for role in request.user.profile.roles.all()]

    if string_any('Editor' in role.name for role in roles):
        viewable = True
        editor = True
        if proposal.requestor and not proposal.requestor == request.user and not 'press-editor' in user_roles:
            editor = False
    else:
        editor = False
   

    template = "core/proposals/assign/view_proposal.html"
    context = {
        'proposal_form': proposal_form,
        'default_fields': default_fields,
        'proposal':proposal,
        'not_readonly':False,
        'data':data,
        'revise':True,
        'assign':True,
        'editor': editor,
        'authors': authors,
        'viewable':viewable,
        'core_proposal':models.ProposalForm.objects.get(pk=proposal_form_id),
    }

    return render(request, template, context)

@is_editor
def proposal_assign_edit(request, proposal_id):

    proposal = submission_models.Proposal.objects.get(pk=proposal_id)
    proposal_form_id = models.Setting.objects.get(name='proposal_form').value

    if proposal.owner == request.user:
        viewable = True

    proposal_form = manager_forms.GeneratedForm(form=models.ProposalForm.objects.get(pk=proposal.form.id))
    default_fields = manager_forms.DefaultForm(initial={'title': proposal.title,'author':proposal.author,'subtitle':proposal.subtitle})

    intial_data={}
    data = {}
    if proposal.data:
        data = json.loads(proposal.data)
        for k,v in data.items():
            intial_data[k] = v[0]


    proposal_form.initial=intial_data
    
    roles = request.user.profile.roles.all()

    user_roles = [role.slug for role in request.user.profile.roles.all()]

    if string_any('Editor' in role.name for role in roles):
        viewable = True
        editor = True
        if proposal.requestor and not proposal.requestor == request.user and not 'press-editor' in user_roles:
            editor = False

        print editor
    else:
        editor = False

    if request.POST and editor:
        proposal_form = manager_forms.GeneratedForm(request.POST, request.FILES, form=models.ProposalForm.objects.get(pk=proposal.form.id))
        default_fields = manager_forms.DefaultForm(request.POST)
        if proposal_form.is_valid() and default_fields.is_valid():

            save_dict = {}
            file_fields = models.ProposalFormElementsRelationship.objects.filter(form=models.ProposalForm.objects.get(pk=proposal.form.id), element__field_type='upload')
            data_fields = models.ProposalFormElementsRelationship.objects.filter(~Q(element__field_type='upload'), form=models.ProposalForm.objects.get(pk=proposal.form.id))

            for field in file_fields:
                if field.element.name in request.FILES:
                    # TODO change value from string to list [value, value_type]
                    save_dict[field.element.name] = [handle_proposal_file_form(request.FILES[field.element.name], proposal, 'other', request.user, "Attachment: Uploaded by %s" % (request.user.username))]

            for field in data_fields:
                if field.element.name in request.POST:
                    # TODO change value from string to list [value, value_type]
                    save_dict[field.element.name] = [request.POST.get(field.element.name), 'text']

            json_data = smart_text(json.dumps(save_dict))
            proposal = submission_models.Proposal.objects.get(form=models.ProposalForm.objects.get(pk=proposal.form.id),pk=proposal_id)
            proposal.data=json_data
            proposal.status = "submission"
            defaults=default_fields.cleaned_data
            proposal.title = defaults.get("title")
            proposal.author = defaults.get("author")
            proposal.subtitle = defaults.get("subtitle")    
            proposal.save()

            update_email_text = models.Setting.objects.get(group__name='email', name='proposal_update_ack').value
            log.add_proposal_log_entry(proposal=proposal,user=request.user, kind='proposal', message='Unassigned Proposal "%s %s" has been updated.'%(proposal.title,proposal.subtitle), short_name='Unassigned Proposal Updated')
            #logic.send_proposal_update(proposal, email_text=update_email_text, sender=request.user, receiver=proposal.owner)
            messages.add_message(request, messages.SUCCESS, 'Unassigned Proposal %s updated' % proposal.id)
            return redirect(reverse('proposals'))

    template = "core/proposals/assign/view_proposal.html"
    context = {
        'proposal_form': proposal_form,
        'default_fields': default_fields,
        'proposal':proposal,
        'not_readonly':True,
        'data':data,
        'revise':True,
        'editor': editor,
        'viewable':viewable,
        'core_proposal':models.ProposalForm.objects.get(pk=proposal_form_id),
    }

    return render(request, template, context)
@is_editor
def proposal(request, user_id = None):
    proposal_list = submission_models.Proposal.objects.filter((~Q(status='declined') & ~Q(status='accepted') & Q(owner__isnull=False)))
    unassigned_proposals = submission_models.Proposal.objects.filter(owner__isnull=True)
   
    proposals = []

    for proposal in proposal_list:
        if user_id:
            if proposal.book_editors.filter(pk=user_id).exists():
                proposals.append(proposal)
        else:
            if 'press-editor' in request.user_roles:
                proposals.append(proposal)
            elif not proposal.requestor:
                proposals.append(proposal)
            elif proposal.requestor==request.user:
                proposals.append(proposal)
            elif proposal.book_editors.filter(username=request.user.username).exists():
                proposals.append(proposal)

    template = 'core/proposals/proposal.html'
    context = {
        'proposal_list': proposals,
        'unassigned_proposal_list': unassigned_proposals,
        'open': True,
        'user_id':user_id,
    }

    return render(request, template, context)

@is_editor
def proposal_history(request):
    proposal_list = submission_models.Proposal.objects.all()
  
    proposals = []
    user_roles = [role.slug for role in request.user.profile.roles.all()]

    for proposal in proposal_list:
        if 'press-editor' in user_roles:
            proposals.append(proposal)
        elif not proposal.requestor:
            proposals.append(proposal)
        elif proposal.requestor==request.user:
            proposals.append(proposal)

    template = 'core/proposals/proposal.html'
    context = {
        'proposal_list': proposals,
    }

    return render(request, template, context)
    
@is_editor
def view_proposal(request, proposal_id):
    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    relationships = models.ProposalFormElementsRelationship.objects.filter(form=proposal.form)
    if proposal.data:
        data = json.loads(proposal.data)
    else:
        data = {}
    
    if not request.POST and request.GET.get('download') == 'docx':
        path = create_proposal_form(proposal)
        return serve_proposal_file(request, path)

    template = 'core/proposals/view_proposal.html'
    context = {
        'proposal': proposal,
        'relationships':relationships,
        'data':data,
    }

    return render(request, template, context)

def create_proposal_form(proposal):
    document = Document()
    document.add_heading(proposal.title, 0)
    p = document.add_paragraph('You should complete this form and then use the proposal page to upload it.')
    relations = models.ProposalFormElementsRelationship.objects.filter(form=proposal.form).order_by('order')
    document.add_heading("Title", level=1)
    document.add_paragraph(proposal.title).italic = True
    document.add_heading("Subtitle", level=1)
    document.add_paragraph(proposal.subtitle).italic = True
    document.add_heading("Author", level=1)
    document.add_paragraph(proposal.author).italic = True

    data = json.loads(proposal.data)
   
    for relation in relations:
        v = data[relation.element.name]
        document.add_heading(relation.element.name, level=1)        
        text = BeautifulSoup(smart_text(v[0]),"html.parser").get_text()
        document.add_paragraph(text).bold = True  

    document.add_page_break()
    if not os.path.exists(os.path.join(settings.BASE_DIR, 'files', 'forms')):
        os.makedirs(os.path.join(settings.BASE_DIR, 'files', 'forms'))
    path = os.path.join(settings.BASE_DIR, 'files', 'forms', '%s.docx' % str(uuid4()))

    document.save(path)
    return path

@is_editor
def withdraw_proposal_review(request, proposal_id,review_id):
    submission = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    review_assignment = get_object_or_404(submission_models.ProposalReview, pk=review_id)
    if review_assignment.withdrawn:
        review_assignment.withdrawn = False
    else:
        review_assignment.withdrawn = True
    review_assignment.save()

    return redirect(reverse('view_proposal', kwargs={'proposal_id': proposal_id}))

@is_editor
def remove_proposal_review(request, proposal_id,review_id):
    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    review_assignment = get_object_or_404(submission_models.ProposalReview, pk=review_id)
    review_assignment.delete()
    review_assignments = submission_models.ProposalReview.objects.filter(proposal=proposal)
    if not review_assignments:
        proposal.review_form = None
        proposal.save()

    return redirect(reverse('view_proposal', kwargs={'proposal_id': proposal_id}))

@is_editor
def start_proposal_review(request, proposal_id):
    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id, date_review_started__isnull=True)
    reviewers = models.User.objects.filter(profile__roles__slug='reviewer')
    committees = manager_models.Group.objects.filter(group_type='review_committee')
    email_text = models.Setting.objects.get(group__name='email', name='proposal_review_request').value
    start_form = submission_forms.ProposalStart()

    if request.POST:
        start_form = submission_forms.ProposalStart(request.POST, request.FILES, instance=proposal)
        if start_form.is_valid():
            if request.FILES.get('attachment'):
                attachment = handle_proposal_file(request.FILES.get('attachment'), proposal, 'misc', request.user)
            else:
                attachment = None
                blind = request.POST.get('blind')
            proposal = start_form.save(commit=False)
            proposal.date_review_started = timezone.now()
            due_date = request.POST.get('due_date')
            email_text = smart_text(request.POST.get('email_text'))

            reviewers = User.objects.filter(pk__in=request.POST.getlist('reviewer'))
            committees = manager_models.Group.objects.filter(pk__in=request.POST.getlist('committee'))

            # Handle reviewers
            for reviewer in reviewers:
                new_review_assignment = submission_models.ProposalReview(
                    user=reviewer,
                    proposal=proposal,
                    due=due_date,
                    blind = blind,
                    requestor = request.user
                )

                try:
                    new_review_assignment.save()
                    proposal.review_assignments.add(new_review_assignment)
                    logic.send_proposal_review_request(proposal, new_review_assignment, email_text, attachment)
                except IntegrityError:
                    messages.add_message(request, messages.WARNING, '%s %s is already a reviewer' % (reviewer.first_name, reviewer.last_name))

            # Handle committees
            for committee in committees:
                members = manager_models.GroupMembership.objects.filter(group=committee)
                for member in members:
                    new_review_assignment = submission_models.ProposalReview(
                        user=member.user,
                        proposal=proposal,
                        due=due_date,
                        blind = blind,
                        requestor = request.user
                    )

                    try:
                        new_review_assignment.save()
                        proposal.review_assignments.add(new_review_assignment)
                        logic.send_proposal_review_request(proposal, new_review_assignment, email_text, attachment = attachment)
                    except IntegrityError:
                        messages.add_message(request, messages.WARNING, '%s %s is already a reviewer' % (member.user.first_name, member.user.last_name))

            # Tidy up and save
            proposal.requestor = request.user
            proposal.date_review_started = timezone.now()
            proposal.save()

            return redirect(reverse('view_proposal', kwargs={'proposal_id': proposal.id}))

    template = 'core/proposals/start_proposal_review.html'
    context = {
        'proposal': proposal,
        'start_form': start_form,
        'reviewers': reviewers,
        'committees': committees,
        'email_text':email_text,
    }

    return render(request, template, context)


@is_reviewer
def view_proposal_review_decision(request, proposal_id, assignment_id):

    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    proposal_form = manager_forms.GeneratedForm(form=models.ProposalForm.objects.get(pk=proposal.form.id))
    default_fields = manager_forms.DefaultForm(initial={'title': proposal.title,'author':proposal.author,'subtitle':proposal.subtitle})
    relationships = models.ProposalFormElementsRelationship.objects.filter(form=proposal.form)
    data = json.loads(proposal.data)
    intial_data={}
    for k,v in data.items():
        intial_data[k] = v[0]

    proposal_form.initial=intial_data
    review_assignment = get_object_or_404(submission_models.ProposalReview, pk=assignment_id, withdrawn = False)
    
    if request.POST:
        if 'accept' in request.POST:
            review_assignment.accepted = timezone.now()
            review_assignment.save()
            message = "Review Assignment request for proposal '%s' has been accepted by %s %s."  % (proposal.title,review_assignment.user.first_name, review_assignment.user.last_name)
            if proposal.requestor:
                notification = models.Task(assignee=proposal.requestor,creator=request.user,text=message,workflow='proposal')
                notification.save()
            else:
                editors = User.objects.filter(profile__roles__slug='press-editor')
                for editor in editors:
                    notification = models.Task(assignee=editor,creator=request.user,text=message,workflow='proposal')
                    notification.save()
            return redirect(reverse('view_proposal_review', kwargs={'proposal_id': proposal.id,'assignment_id': assignment_id}))
                
        elif 'decline' in request.POST:
            review_assignment.declined = timezone.now()
            review_assignment.save()
            message = "Review Assignment request for proposal '%s' has been declined by %s %s."  % (proposal.title,review_assignment.user.first_name, review_assignment.user.last_name)
            if proposal.requestor:
                notification = models.Task(assignee=proposal.requestor,creator=request.user,text=message,workflow='proposal')
                notification.save()
            else:
                editors = User.objects.filter(profile__roles__slug='press-editor')
                for editor in editors:
                    notification = models.Task(assignee=editor,creator=request.user,text=message,workflow='proposal')
                    notification.save()
            return redirect(reverse('view_proposal_review', kwargs={'proposal_id': proposal.id,'assignment_id': assignment_id}))


    
    template = 'core/proposals/decision_review_assignment.html'
    context = {
        'proposal': proposal,
        'proposal_form':proposal_form,
        'review': review_assignment,
        'data': data,
        'active': 'proposal_review',
        'relationships':relationships,
        'instructions': models.Setting.objects.get(group__name='general', name='instructions_for_task_proposal').value
    }

    return render(request, template, context)

@is_reviewer
def view_completed_proposal_review(request, proposal_id, assignment_id):

    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    proposal_form = manager_forms.GeneratedForm(form=models.ProposalForm.objects.get(pk=proposal.form.id))
    default_fields = manager_forms.DefaultForm(initial={'title': proposal.title,'author':proposal.author,'subtitle':proposal.subtitle})
    relationships = models.ProposalFormElementsRelationship.objects.filter(form=proposal.form)
    data = json.loads(proposal.data)
    intial_data={}
    for k,v in data.items():
        intial_data[k] = v[0]

    proposal_form.initial = intial_data
    review_assignment = get_object_or_404(submission_models.ProposalReview, pk=assignment_id, withdrawn = False)
    result = review_assignment.results
    form = review_forms.GeneratedForm(form=proposal.review_form)

    ci_required = models.Setting.objects.get(group__name='general', name='ci_required')
    recommendation_form = forms.RecommendationForm(ci_required=ci_required.value)
    if result:
        relations = review_models.FormElementsRelationship.objects.filter(form=result.form)
        data_ordered = logic.order_data(logic.decode_json(result.data), relations)
    else:
        relations = None
        data_ordered = None

    if not request.POST and request.GET.get('download') == 'docx':
        path = create_proposal_review_form(review_assignment)
        return serve_proposal_file(request, path)
    elif request.POST:
        form = review_forms.GeneratedForm(request.POST, request.FILES, form=proposal.review_form)
        recommendation_form = forms.RecommendationForm(request.POST, ci_required=ci_required.value)
        if form.is_valid() and recommendation_form.is_valid():
            save_dict = {}
            file_fields = review_models.FormElementsRelationship.objects.filter(form=proposal.review_form, element__field_type='upload')
            data_fields = review_models.FormElementsRelationship.objects.filter(~Q(element__field_type='upload'), form=proposal.review_form)

            for field in file_fields:
                if field.element.name in request.FILES:
                    # TODO change value from string to list [value, value_type]
                    save_dict[field.element.name] = [handle_review_file(request.FILES[field.element.name], submission, review_assignment, 'reviewer')]

            for field in data_fields:
                if field.element.name in request.POST:
                    # TODO change value from string to list [value, value_type]
                    save_dict[field.element.name] = [request.POST.get(field.element.name), 'text']

            json_data = smart_text(json.dumps(save_dict))
            form_results = review_models.FormResult(form=proposal.review_form, data=json_data)
            form_results.save()

            #if request.FILES.get('review_file_upload'):
            #   handle_review_file(request.FILES.get('review_file_upload'), submission, review_assignment, 'reviewer')

            review_assignment.completed = timezone.now()
            if not review_assignment.accepted:
                review_assignment.accepted = timezone.now()
            review_assignment.recommendation = request.POST.get('recommendation')
            review_assignment.competing_interests = request.POST.get('competing_interests')
            review_assignment.results = form_results
            review_assignment.save()

            return redirect(reverse('user_dashboard'))



    template = 'core/proposals/completed_review_assignment.html'
    context = {
        'proposal': proposal,
        'proposal_form':proposal_form,
        'review_assignment': review_assignment,
        'data_ordered': data_ordered,
        'data_ordered_size': len(data_ordered),
        'result': result,
        'form':form,
        'recommendation_form':recommendation_form,
        'active': 'proposal_review',
        'relationships':relationships,
        'instructions': models.Setting.objects.get(group__name='general', name='instructions_for_task_proposal').value,
        'data': data,
    }

    return render(request, template, context)

def get_list_of_editors(proposal):
    book_editors = proposal.book_editors.all()
    previous_editors=[]
    for book_editor in book_editors:
        previous_editors.append(book_editor)
    all_book_editors = User.objects.filter(profile__roles__slug='book-editor')
    
    list_of_editors =[{} for t in range(0,len(all_book_editors)) ]  
    for t,editor in enumerate(all_book_editors):
        already_added = False
        if editor in previous_editors:
            already_added = True
        list_of_editors[t] = {'editor': editor, 'already_added': already_added,}
    return list_of_editors


@is_press_editor
def proposal_add_editors(request, proposal_id):
    
    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    
    list_of_editors = get_list_of_editors(proposal)

    email_text = models.Setting.objects.get(group__name='email', name='book_editor_ack').value
    
    if request.GET and "add" in request.GET:
        user_id = request.GET.get("add")
        user = User.objects.get(pk=user_id)
        proposal.book_editors.add(user)
        proposal.save()
        list_of_editors = get_list_of_editors(proposal)

    elif request.GET and "remove" in request.GET:
        user_id = request.GET.get("remove")
        user = User.objects.get(pk=user_id)
        proposal.book_editors.remove(user)
        proposal.save()
        list_of_editors = get_list_of_editors(proposal)




    template = 'core/proposals/add_editors.html'
    context = {
        'proposal': proposal,
        'list_of_editors':list_of_editors,
    }

    return render(request, template, context)


@is_reviewer
def view_proposal_review(request, proposal_id, assignment_id):

    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    proposal_form = manager_forms.GeneratedForm(form=models.ProposalForm.objects.get(pk=proposal.form.id))
    default_fields = manager_forms.DefaultForm(initial={'title': proposal.title,'author':proposal.author,'subtitle':proposal.subtitle})
    relationships = models.ProposalFormElementsRelationship.objects.filter(form=proposal.form)
    data = json.loads(proposal.data)

    intial_data = {}
    for k,v in data.items():
        intial_data[k] = v[0]

    proposal_form.initial=intial_data
    review_assignment = get_object_or_404(submission_models.ProposalReview, pk=assignment_id, withdrawn = False)
    result = review_assignment.results

    form = review_forms.GeneratedForm(form=proposal.review_form)
    
    if review_assignment.reopened:
        result = review_assignment.results
        if result:
            initial_data = {}
            data = json.loads(result.data)
            for k,v in data.items():
                initial_data[k] = v[0]
            form.initial = initial_data

    ci_required = models.Setting.objects.get(group__name='general', name='ci_required')
    recommendation_form = forms.RecommendationForm(ci_required=ci_required.value)
    
    if review_assignment.reopened:
        initial_data = {}
        initial_data[u'recommendation']=review_assignment.recommendation
        initial_data[u'competing_interests']=review_assignment.competing_interests
        recommendation_form.initial=initial_data

    if result:
        relations = review_models.FormElementsRelationship.objects.filter(form=result.form)
        data_ordered = logic.order_data(logic.decode_json(result.data), relations)
    else:
        relations = None
        data_ordered = None

    if not request.POST and request.GET.get('download') == 'docx':
        path = create_proposal_review_form(review_assignment)
        return serve_proposal_file(request, path)
    elif request.POST:
        form = review_forms.GeneratedForm(request.POST, request.FILES, form=proposal.review_form)
        recommendation_form = forms.RecommendationForm(request.POST, ci_required=ci_required.value)
        if form.is_valid() and recommendation_form.is_valid():
            save_dict = {}
            file_fields = review_models.FormElementsRelationship.objects.filter(form=proposal.review_form, element__field_type='upload')
            data_fields = review_models.FormElementsRelationship.objects.filter(~Q(element__field_type='upload'), form=proposal.review_form)

            for field in file_fields:
                if field.element.name in request.FILES:
                    # TODO change value from string to list [value, value_type]
                    save_dict[field.element.name] = [handle_review_file(request.FILES[field.element.name], submission, review_assignment, 'reviewer')]

            for field in data_fields:
                if field.element.name in request.POST:
                    # TODO change value from string to list [value, value_type]
                    save_dict[field.element.name] = [request.POST.get(field.element.name), 'text']

            json_data = smart_text(json.dumps(save_dict))
            form_results = review_models.FormResult(form=proposal.review_form, data=json_data)
            form_results.save()

            #if request.FILES.get('review_file_upload'):
            #   handle_review_file(request.FILES.get('review_file_upload'), submission, review_assignment, 'reviewer')

            review_assignment.completed = timezone.now()
            if not review_assignment.accepted:
                review_assignment.accepted = timezone.now()
            review_assignment.recommendation = request.POST.get('recommendation')
            review_assignment.competing_interests = request.POST.get('competing_interests')
            review_assignment.results = form_results
            review_assignment.reopened = False
            review_assignment.save()
            message = "Review assignment for proposal '%s' has been completed by %s ."  % (review_assignment.proposal.title,review_assignment.user.profile.full_name())
            notification = models.Task(assignee=review_assignment.proposal.requestor,creator=request.user,text=message,workflow='proposal')
            notification.save()

            return redirect(reverse('user_dashboard'))



    template = 'core/proposals/review_assignment.html'
    context = {
        'proposal': proposal,
        'proposal_form':proposal_form,
        'review_assignment': review_assignment,
        'relationships':relationships,
        'data_ordered': data_ordered,
        'result': result,
        'form': form,
        'recommendation_form':recommendation_form,
        'active': 'proposal_review',
        'instructions': models.Setting.objects.get(group__name='general', name='instructions_for_task_proposal').value,
        'data': data,
    }

    return render(request, template, context)

@is_editor
def add_proposal_reviewers(request, proposal_id):

    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    reviewers = models.User.objects.filter(profile__roles__slug='reviewer')
    committees = manager_models.Group.objects.filter(group_type='review_committee')
    email_text = models.Setting.objects.get(group__name='email', name='proposal_review_request').value
    start_form = submission_forms.ProposalStart()
    
    if request.POST:
        if not proposal.review_form:
            start_form = submission_forms.ProposalStart(request.POST, request.FILES, instance=proposal)
            if start_form.is_valid():
                proposal = start_form.save(commit=False)
                proposal.save()
        if request.FILES.get('attachment'):
            attachment = handle_proposal_file(request.FILES.get('attachment'), proposal, 'misc', request.user)
        else:
            attachment = None
        print request.FILES
        due_date = request.POST.get('due_date')
        blind = request.POST.get('blind')
        reviewers = User.objects.filter(pk__in=request.POST.getlist('reviewer'))
        committees = manager_models.Group.objects.filter(pk__in=request.POST.getlist('committee'))
        email_text = request.POST.get('email_text')

        # Handle reviewers
        for reviewer in reviewers:
            new_review_assignment = submission_models.ProposalReview(
                user=reviewer,
                proposal=proposal,
                due=due_date,
                blind = blind,
                requestor = request.user
            )

            try:
                new_review_assignment.save()
                proposal.review_assignments.add(new_review_assignment)
                logic.send_proposal_review_request(proposal, new_review_assignment, email_text, attachment)
            except IntegrityError:
                messages.add_message(request, messages.WARNING, '%s %s is already a reviewer' % (reviewer.first_name, reviewer.last_name))

        # Handle committees
        for committee in committees:
            members = manager_models.GroupMembership.objects.filter(group=committee)
            for member in members:
                new_review_assignment = submission_models.ProposalReview(
                    user=member.user,
                    proposal=proposal,
                    due=due_date,
                    blind = blind,
                    requestor = request.user
                )

                try:
                    new_review_assignment.save()
                    proposal.review_assignments.add(new_review_assignment)
                    logic.send_proposal_review_request(proposal, new_review_assignment, email_text, attachment)
                except IntegrityError:
                    messages.add_message(request, messages.WARNING, '%s %s is already a reviewer' % (member.user.first_name, member.user.last_name))

        # Tidy up and save

        proposal.date_review_started = timezone.now()
        proposal.save()

        return redirect(reverse('view_proposal', kwargs={'proposal_id': proposal.id}))

    template = 'core/proposals/add_reviewers.html'
    context = {
        'proposal': proposal,
        'reviewers': reviewers,
        'committees': committees,
        'email_text':email_text,
        'start_form':start_form,
    }

    return render(request, template, context)

@is_editor
def decline_proposal(request, proposal_id):

    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    email_text = models.Setting.objects.get(group__name='email', name='proposal_decline')

    if request.POST:
        proposal.status = 'declined'
        logic.close_active_reviews(proposal)
        proposal.requestor=request.user
        proposal.save()
        log.add_proposal_log_entry(proposal=proposal,user=request.user, kind='proposal', message='Proposal "%s %s" was declined.'%(proposal.title,proposal.subtitle), short_name='Proposal Declined')
        logic.send_proposal_decline(proposal, email_text=request.POST.get('decline-email'), sender=request.user)
        return redirect(reverse('proposals'))

    template = 'core/proposals/decline_proposal.html'
    context = {
        'proposal': proposal,
        'email_text': logic.setting_template_loader(setting=email_text,path="core/proposals/",pattern="email_template_",dictionary={'sender':request.user,'receiver':proposal.owner,'proposal':proposal,
        'press_name':models.Setting.objects.get(group__name='general', name='press_name').value}),
    }

    return render(request, template, context)


@is_editor
def accept_proposal(request, proposal_id):
    'Marks a proposal as accepted, creates a submission and emails the user'
    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    email_text = models.Setting.objects.get(group__name='email', name='proposal_accept')

    if request.POST:
        proposal.status = 'accepted'
        logic.close_active_reviews(proposal)
        proposal.requestor=request.user
        submission = logic.create_submission_from_proposal(proposal, proposal_type=proposal.book_type)
        submission.proposal = proposal
        submission.save()
        attachment = handle_attachment(request, submission)
        logic.send_proposal_accept(proposal, email_text=request.POST.get('accept-email'), submission=submission, sender=request.user, attachment=attachment)
        proposal.date_accepted = timezone.now()
        proposal.save()
        log.add_proposal_log_entry(proposal=proposal,user=request.user, kind='proposal',  message='Proposal "%s %s" was accepted.'%(proposal.title,proposal.subtitle), short_name='Proposal Accepted')
    
        return redirect(reverse('proposals'))

    template = 'core/proposals/accept_proposal.html'
    
    context = {
        'proposal': proposal,
        'email_text': logic.setting_template_loader(setting=email_text,path="core/proposals/",pattern="email_template_",dictionary={'sender':request.user,'receiver':proposal.owner,'proposal':proposal,
        'press_name':models.Setting.objects.get(group__name='general', name='press_name').value}),
    }

    return render(request, template, context)

@is_editor
def reopen_proposal_review(request, proposal_id, assignment_id):

    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    review_assignment = get_object_or_404(submission_models.ProposalReview, pk=assignment_id)
    email_text = models.Setting.objects.get(group__name='email', name='proposal_review_request')

    if request.POST:
        review_assignment.reopened = True
        review_assignment.comments_from_editor = request.POST.get('comments')
        review_assignment.completed = None
        review_assignment.save()
        email_updated_text =request.POST.get('email')
        due_date =request.POST.get('due_date') 
        review_assignment.due = datetime.strptime(due_date, "%Y-%m-%d")
        review_assignment.save()
        email_updated_text = string.replace(email_updated_text,'_due_date_',due_date)
        logic.send_proposal_review_request(proposal, review_assignment, email_updated_text)
        log.add_proposal_log_entry(proposal=proposal,user=request.user, kind='proposal', message='Revisions request for proposal %s %s.'%(proposal.title,proposal.subtitle), short_name='Proposal Revisions Requested')
    
        return redirect(reverse('proposals'))

    template = 'core/proposals/reopen_review.html'
    context = {
        'proposal': proposal,
        'email_text': logic.setting_template_loader(setting=email_text,path="core/proposals/",pattern="email_template_",dictionary={'sender':request.user,'receiver':proposal.owner,'proposal':proposal,
        'press_name':models.Setting.objects.get(group__name='general', name='press_name').value}),
    }

    return render(request, template, context)

@is_editor
def request_proposal_revisions(request, proposal_id):

    proposal = get_object_or_404(submission_models.Proposal, pk=proposal_id)
    email_text = models.Setting.objects.get(group__name='email', name='proposal_request_revisions')

    if request.POST:
        proposal.status = 'revisions_required'
        logic.close_active_reviews(proposal)
        proposal.requestor=request.user
        email_updated_text =request.POST.get('revisions-email')
        due_date =request.POST.get('due_date') 
        proposal.revision_due_date = datetime.strptime(due_date, "%Y-%m-%d")
        proposal.save()
        email_updated_text = string.replace(email_updated_text,'_due_date_',due_date)
        logic.send_proposal_revisions(proposal, email_text=email_updated_text, sender=request.user)
        
        log.add_proposal_log_entry(proposal=proposal,user=request.user, kind='proposal', message='Revisions request for proposal %s %s.'%(proposal.title,proposal.subtitle), short_name='Proposal Revisions Requested')
    
        return redirect(reverse('proposals'))

    template = 'core/proposals/revisions_proposal.html'
    context = {
        'proposal': proposal,
        'email_text': logic.setting_template_loader(setting=email_text,path="core/proposals/",pattern="email_template_",dictionary={'sender':request.user,'receiver':proposal.owner,'proposal':proposal,
        'press_name':models.Setting.objects.get(group__name='general', name='press_name').value}),
    }

    return render(request, template, context)

def render_choices(choices):
    c_split = choices.split('|')
    return [(choice.capitalize(), choice) for choice in c_split]

@is_reviewer
def create_proposal_review_form(proposal):
    document = Document()
    document.add_heading(proposal.proposal.title, 0)
    p = document.add_paragraph('You should complete this form and then use the review page to upload it.')
    relations = review_models.FormElementsRelationship.objects.filter(form=proposal.proposal.review_form).order_by('order')
    for relation in relations:

        if relation.element.field_type in ['text', 'textarea', 'date', 'email']:
            document.add_heading(relation.element.name+": _______________________________", level=1)
            document.add_paragraph(relation.help_text).italic = True

        if relation.element.field_type in ['select', 'check']:
            document.add_heading(relation.element.name, level=1)
            if relation.element.field_type == 'select':
                choices = render_choices(relation.element.choices)
            else:
                choices = ['Y', 'N']

            p = document.add_paragraph(relation.help_text)
            p.add_run(' Mark your choice however you like, as long as it is clear.').italic = True
            table = document.add_table(rows=2, cols=len(choices))
            hdr_cells = table.rows[0].cells
            for i, choice in enumerate(choices):
                hdr_cells[i].text = choice[0]
            table.style = 'TableGrid'

    document.add_page_break()
    if not os.path.exists(os.path.join(settings.BASE_DIR, 'files', 'forms')):
        os.makedirs(os.path.join(settings.BASE_DIR, 'files', 'forms'))
    path = os.path.join(settings.BASE_DIR, 'files', 'forms', '%s.docx' % str(uuid4()))

    document.save(path)
    return path

@is_reviewer
def serve_proposal_file(request, file_path):
    try:
        fsock = open(file_path, 'r')
        mimetype = mimetypes.guess_type(file_path)
        response = StreamingHttpResponse(fsock, content_type=mimetype)
        response['Content-Disposition'] = "attachment; filename=proposal_form.docx"
        
        return response
    except IOError:
        messages.add_message(request, messages.ERROR, 'File not found.')
        return HttpResponseRedirect(request.META.get('HTTP_REFERER'))
        
## END PROPOSAL ##

